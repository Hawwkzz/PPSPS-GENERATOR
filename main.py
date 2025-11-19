# app/main.py — FREEFORM PPSPS (Markdown → DOCX), no placeholders
from __future__ import annotations

# ===== Stdlib =====
import os, glob, tempfile, re, json
from io import BytesIO
from pathlib import Path
from urllib.parse import unquote
from datetime import datetime, date
from typing import Optional, Literal
import csv
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
import base64
import mimetypes
import logging
logger = logging.getLogger("sps")
if not logger.handlers:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

# ===== PLACEHOLDERS FORMULAIRE =====
FORM_PLACEHOLDERS = {
    "{{NOM_PROJET}}": "",
    "{{ADRESSE_CHANTIER}}": "",
    "{{REFERENCE_AFFAIRE}}": "",
    "{{TELEPHONE_CHANTIER}}": "",
    "{{DUREE_SEMAINES}}": "",
    "{{EFFECTIF_MAXIMUM}}": "",
    "{{LOTS_TRAVAUX}}": "",
    "{{NOM_ENTREPRISE}}": "",
    "{{ADRESSE_ENTREPRISE}}": "",
    "{{TELEPHONE_ENTREPRISE}}": "",
    "{{EMAIL_ENTREPRISE}}": "",
    "{{RESPONSABLE_TRAVAUX}}": "",
    "{{MAITRE_OUVRAGE}}": "",
    "{{MAITRE_OEUVRE}}": ""
}



# ===== FastAPI / Starlette =====
from fastapi import FastAPI, Body, Request, Form, Depends, HTTPException, UploadFile, File
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse, FileResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.sessions import SessionMiddleware
from starlette.datastructures import FormData

# ===== DB / Models =====
from sqlmodel import Session, select
from passlib.context import CryptContext
from pydantic import BaseModel, Field as PydField, field_validator
from app.db import init_db, get_session, engine

from app.config import (
    APP_TITLE, SECRET_KEY, ADMIN_TOKEN,
)
from app.db import init_db, get_session
from app.models import ProjectDB, DocumentDB, AttachmentDB, UserDB
# ===== Modèles Tokens =====
from sqlalchemy import ForeignKey
from sqlalchemy.orm import Mapped, mapped_column
from app.models import Base

class TokenPackage(Base):
    """Package de jetons (1 jeton = 50€)"""
    __tablename__ = "token_package"
    id: Mapped[int] = mapped_column(primary_key=True)
    name: Mapped[str]
    tokens: Mapped[int]
    price_cents: Mapped[int]
    stripe_price_id: Mapped[Optional[str]] = mapped_column(nullable=True)
    active: Mapped[bool] = mapped_column(default=True)
    created_at: Mapped[datetime] = mapped_column(default=datetime.utcnow)

class UserTokenBalance(Base):
    """Solde de jetons par utilisateur"""
    __tablename__ = "user_token_balance"
    id: Mapped[int] = mapped_column(primary_key=True)
    user_id: Mapped[int] = mapped_column(ForeignKey("userdb.id"), unique=True, index=True)
    balance: Mapped[int] = mapped_column(default=0)
    total_purchased: Mapped[int] = mapped_column(default=0)
    total_used: Mapped[int] = mapped_column(default=0)
    updated_at: Mapped[datetime] = mapped_column(default=datetime.utcnow)

class TokenTransaction(Base):
    """Historique des transactions"""
    __tablename__ = "token_transaction"
    id: Mapped[int] = mapped_column(primary_key=True)
    user_id: Mapped[int] = mapped_column(ForeignKey("userdb.id"), index=True)
    type: Mapped[str]  # "purchase", "usage", "refund"
    amount: Mapped[int]
    balance_after: Mapped[int]
    description: Mapped[str]
    payment_intent_id: Mapped[Optional[str]] = mapped_column(nullable=True)
    package_id: Mapped[Optional[int]] = mapped_column(ForeignKey("token_package.id"), nullable=True)
    project_id: Mapped[Optional[int]] = mapped_column(ForeignKey("projectdb.id"), nullable=True)
    document_id: Mapped[Optional[int]] = mapped_column(ForeignKey("documentdb.id"), nullable=True)
    created_at: Mapped[datetime] = mapped_column(default=datetime.utcnow)

class StripeWebhookEvent(Base):
    """Log webhooks Stripe"""
    __tablename__ = "stripe_webhook_event"
    id: Mapped[int] = mapped_column(primary_key=True)
    stripe_event_id: Mapped[str] = mapped_column(unique=True, index=True)
    event_type: Mapped[str]
    processed: Mapped[bool] = mapped_column(default=False)
    error: Mapped[Optional[str]] = mapped_column(nullable=True)
    created_at: Mapped[datetime] = mapped_column(default=datetime.utcnow)

    # ===== Services Tokens & Stripe =====
class InsufficientTokensError(Exception):
    """Levée quand pas assez de jetons"""
    pass

class TokenService:
    """Service de gestion des jetons"""
    
    @staticmethod
    def get_or_create_balance(session: Session, user_id: int) -> UserTokenBalance:
        balance = session.exec(
            select(UserTokenBalance).where(UserTokenBalance.user_id == user_id)
        ).first()
        if not balance:
            balance = UserTokenBalance(user_id=user_id, balance=0)
            session.add(balance)
            session.commit()
            session.refresh(balance)
        return balance
    
    @staticmethod
    def get_balance(session: Session, user_id: int) -> int:
        balance = TokenService.get_or_create_balance(session, user_id)
        return balance.balance
    
    @staticmethod
    def add_tokens(session: Session, user_id: int, amount: int, description: str,
                   payment_intent_id: Optional[str] = None, package_id: Optional[int] = None):
        balance = TokenService.get_or_create_balance(session, user_id)
        balance.balance += amount
        balance.total_purchased += amount
        balance.updated_at = datetime.utcnow()
        
        tx = TokenTransaction(
            user_id=user_id, type="purchase", amount=amount,
            balance_after=balance.balance, description=description,
            payment_intent_id=payment_intent_id, package_id=package_id
        )
        session.add(balance)
        session.add(tx)
        session.commit()
        return tx
    
    @staticmethod
    def use_token(session: Session, user_id: int, project_id: int,
                  document_id: Optional[int] = None, description: str = "Génération PPSPS"):
        balance = TokenService.get_or_create_balance(session, user_id)
        if balance.balance < 1:
            raise InsufficientTokensError(
                f"Solde insuffisant. Vous avez {balance.balance} jeton(s), 1 jeton requis."
            )
        balance.balance -= 1
        balance.total_used += 1
        balance.updated_at = datetime.utcnow()
        
        tx = TokenTransaction(
            user_id=user_id, type="usage", amount=-1,
            balance_after=balance.balance, description=description,
            project_id=project_id, document_id=document_id
        )
        session.add(balance)
        session.add(tx)
        session.commit()
        return tx
    
    @staticmethod
    def refund_token(session: Session, user_id: int, project_id: int, reason: str = "Remboursement"):
        balance = TokenService.get_or_create_balance(session, user_id)
        balance.balance += 1
        balance.total_used -= 1
        balance.updated_at = datetime.utcnow()
        
        tx = TokenTransaction(
            user_id=user_id, type="refund", amount=1,
            balance_after=balance.balance, description=reason, project_id=project_id
        )
        session.add(balance)
        session.add(tx)
        session.commit()
        return tx
    
    @staticmethod
    def get_transactions(session: Session, user_id: int, limit: int = 50):
        return session.exec(
            select(TokenTransaction)
            .where(TokenTransaction.user_id == user_id)
            .order_by(TokenTransaction.created_at.desc())
            .limit(limit)
        ).all()

class StripeService:
    """Service paiement Stripe"""
    
    @staticmethod
    def create_checkout_session(user_id: int, package_id: int, success_url: str,
                                cancel_url: str, session_db: Session):
        package = session_db.get(TokenPackage, package_id)
        if not package or not package.active:
            raise ValueError("Package invalide")
        
        checkout_session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[{
                'price_data': {
                    'currency': 'eur',
                    'unit_amount': package.price_cents,
                    'product_data': {
                        'name': package.name,
                        'description': f"{package.tokens} jeton pour générer un PPSPS",
                    },
                },
                'quantity': 1,
            }],
            mode='payment',
            success_url=success_url,
            cancel_url=cancel_url,
            metadata={'user_id': str(user_id), 'package_id': str(package_id), 'tokens': str(package.tokens)}
        )
        return checkout_session
    
    @staticmethod
    def handle_webhook(payload: bytes, sig_header: str, session_db: Session):
        try:
            event = stripe.Webhook.construct_event(payload, sig_header, STRIPE_WEBHOOK_SECRET)
        except Exception:
            return {'success': False, 'message': 'Signature invalide'}
        
        # Vérifier si déjà traité
        existing = session_db.exec(
            select(StripeWebhookEvent).where(StripeWebhookEvent.stripe_event_id == event['id'])
        ).first()
        if existing and existing.processed:
            return {'success': True, 'message': 'Déjà traité'}
        
        # Logger
        log = StripeWebhookEvent(stripe_event_id=event['id'], event_type=event['type'], processed=False)
        session_db.add(log)
        session_db.commit()
        
        try:
            if event['type'] == 'checkout.session.completed':
                session_stripe = event['data']['object']
                user_id = int(session_stripe['metadata']['user_id'])
                package_id = int(session_stripe['metadata']['package_id'])
                tokens = int(session_stripe['metadata']['tokens'])
                payment_intent_id = session_stripe.get('payment_intent')
                
                package = session_db.get(TokenPackage, package_id)
                TokenService.add_tokens(
                    session=session_db, user_id=user_id, amount=tokens,
                    description=f"Achat {package.name}",
                    payment_intent_id=payment_intent_id, package_id=package_id
                )
            
            log.processed = True
            session_db.add(log)
            session_db.commit()
            return {'success': True, 'message': f"Événement {event['type']} traité"}
        except Exception as e:
            log.error = str(e)
            session_db.add(log)
            session_db.commit()
            return {'success': False, 'message': str(e)}
    
    @staticmethod
    def get_packages(session_db: Session):
        return session_db.exec(
            select(TokenPackage).where(TokenPackage.active == True).order_by(TokenPackage.tokens.asc())
        ).all()
    
    @staticmethod
    def create_default_packages(session_db: Session):
        """Crée le package : 1 jeton = 50€"""
        pkg_data = {'name': 'Génération PPSPS', 'tokens': 1, 'price_cents': 5000}
        existing = session_db.exec(
            select(TokenPackage).where(TokenPackage.name == pkg_data['name'])
        ).first()
        if not existing:
            package = TokenPackage(**pkg_data)
            session_db.add(package)
            session_db.commit()

            # ===== Configuration SEO =====
class SEOConfig:
    SITE_NAME = "PPSPS GENERATOR"
    SITE_DOMAIN = "ppsps-generator.fr"  # À MODIFIER avec ton domaine
    SITE_URL = f"https://{SITE_DOMAIN}"
    
    DEFAULT_TITLE = "PPSPS Generator - Génération automatique de PPSPS par IA"
    DEFAULT_DESCRIPTION = (
        "Générez automatiquement jusqu'à 90% de votre PPSPS en quelques minutes. "
        "Téléversez vos documents (PGC, plans), notre IA produit un PPSPS conforme."
    )
    DEFAULT_KEYWORDS = [
        "PPSPS", "génération PPSPS", "PPSPS automatique", "IA BTP",
        "prévention BTP", "sécurité chantier", "coordonnateur SPS"
    ]
    OG_IMAGE = f"{SITE_URL}/static/og-image.jpg"
    TWITTER_HANDLE = "@SPSCopilot"
    
    @staticmethod
    def get_meta_tags(title: Optional[str] = None, description: Optional[str] = None,
                     keywords: Optional[list[str]] = None, canonical_url: Optional[str] = None,
                     og_type: str = "website", no_index: bool = False):
        full_title = f"{title} - {SEOConfig.SITE_NAME}" if title else SEOConfig.DEFAULT_TITLE
        final_desc = description or SEOConfig.DEFAULT_DESCRIPTION
        all_kw = SEOConfig.DEFAULT_KEYWORDS.copy()
        if keywords: all_kw.extend(keywords)
        
        return {
            "title": full_title,
            "description": final_desc,
            "keywords": ", ".join(all_kw),
            "canonical": canonical_url,
            "og_title": full_title,
            "og_description": final_desc,
            "og_type": og_type,
            "og_image": SEOConfig.OG_IMAGE,
            "og_url": canonical_url or SEOConfig.SITE_URL,
            "twitter_card": "summary_large_image",
            "twitter_site": SEOConfig.TWITTER_HANDLE,
            "robots": "noindex, nofollow" if no_index else "index, follow"
        }

SEO_PAGES = {
    "home": {
        "title": "Génération automatique de PPSPS par IA | BTP",
        "description": "Générez automatiquement jusqu'à 90% de votre PPSPS en quelques minutes avec notre IA spécialisée BTP. Conforme, rapide, économique. Essai gratuit.",
        "keywords": ["génération PPSPS", "PPSPS automatique", "IA BTP", "générateur PPSPS", "PPSPS en ligne", "automatisation BTP"]
    },
    "register": {
        "title": "Créer un compte gratuit | Inscription",
        "description": "Créez votre compte gratuit sur PPSPS Generator et commencez à générer vos PPSPS automatiquement en quelques minutes. Sans engagement.",
        "keywords": ["inscription PPSPS", "créer compte générateur", "essai gratuit BTP"]
    },
    "shop": {
        "title": "Tarifs - Acheter des jetons PPSPS",
        "description": "50€ par génération de PPSPS. Achetez vos jetons et générez des documents conformes en quelques minutes. Paiement sécurisé Stripe.",
        "keywords": ["tarifs PPSPS", "prix génération PPSPS", "acheter jetons", "coût PPSPS"]
    }
}


# === OpenRouter / OpenAI ===
from app.config import OPENROUTER_API_KEY, OPENROUTER_BASE_URL, OPENROUTER_DEFAULT_MODEL
from openai import OpenAI
# === Stripe & Tokens ===
import stripe
from datetime import datetime

# Configuration Stripe
stripe.api_key = os.getenv("STRIPE_SECRET_KEY")
STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET")

client = OpenAI(
    base_url=OPENROUTER_BASE_URL,
    api_key=OPENROUTER_API_KEY,
    default_headers={"HTTP-Referer": "http://localhost:8000", "X-Title": "SPS Copilot"},
)

print("### MODEL IN USE:", OPENROUTER_DEFAULT_MODEL, flush=True)

# ===== Files / text =====
from pypdf import PdfReader
from docx import Document as DocxDocument


# Chemin du template PPSPS (relatif au dossier du projet)
import os as _os
_current_dir = _os.path.dirname(_os.path.abspath(__file__))
TEMPLATE_PATH = _os.path.join(_current_dir, "248234559-modele-PPSPS.docx")

# ===== Template Filler (intégré) =====
"""
Système de remplissage intelligent du template PPSPS.
Utilise le template DOCX original et le remplit via l'IA avec flexibilité.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Dict, Any, List
import re
import json


def replace_placeholders_in_doc(doc, placeholder_values):
    """Remplace tous les placeholders dans un document DOCX de manière robuste."""
    from docx import Document
    
    logger.info("[PLACEHOLDERS] Début du remplacement dans le document")
    replacements_count = 0
    
    def replace_in_paragraph(paragraph, placeholder_values, location=""):
        """Remplace les placeholders dans un paragraphe, même s'ils sont fragmentés."""
        nonlocal replacements_count
        
        for key, value in placeholder_values.items():
            if key in paragraph.text:
                # Reconstruire le texte en fusionnant les runs
                full_text = paragraph.text
                new_text = full_text.replace(key, str(value) if value else "")
                
                # Si le texte a changé, remplacer tout le paragraphe
                if new_text != full_text:
                    logger.info(f"[PLACEHOLDERS] Trouvé '{key}' dans {location}")
                    logger.info(f"  Avant: {full_text[:100]}")
                    logger.info(f"  Après: {new_text[:100]}")
                    replacements_count += 1
                    
                    # Garder le style du premier run
                    first_run_style = paragraph.runs[0].style if paragraph.runs else None
                    first_run_font = paragraph.runs[0].font if paragraph.runs else None
                    
                    # Supprimer tous les runs
                    for _ in range(len(paragraph.runs)):
                        paragraph._element.remove(paragraph.runs[0]._element)
                    
                    # Ajouter un nouveau run avec le texte remplacé
                    new_run = paragraph.add_run(new_text)
                    if first_run_style:
                        new_run.style = first_run_style
    
    # Remplacer dans les paragraphes
    logger.info("[PLACEHOLDERS] Scan des paragraphes...")
    for i, paragraph in enumerate(doc.paragraphs):
        replace_in_paragraph(paragraph, placeholder_values, f"paragraphe {i}")
    
    # Remplacer dans les tableaux
    logger.info("[PLACEHOLDERS] Scan des tableaux...")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, paragraph in enumerate(cell.paragraphs):
                    replace_in_paragraph(paragraph, placeholder_values, f"tableau {t_idx}, ligne {r_idx}, cellule {c_idx}, para {p_idx}")
    
    # Remplacer dans les en-têtes et pieds de page
    logger.info("[PLACEHOLDERS] Scan des en-têtes et pieds de page...")
    for s_idx, section in enumerate(doc.sections):
        for p_idx, paragraph in enumerate(section.header.paragraphs):
            replace_in_paragraph(paragraph, placeholder_values, f"en-tête section {s_idx}, para {p_idx}")
        
        for p_idx, paragraph in enumerate(section.footer.paragraphs):
            replace_in_paragraph(paragraph, placeholder_values, f"pied de page section {s_idx}, para {p_idx}")
    
    logger.info(f"[PLACEHOLDERS] Remplacement terminé: {replacements_count} placeholders remplacés")

class TemplateFiller:
    """Remplit intelligemment le template PPSPS avec les données du projet."""
    
    def __init__(self, template_path: str, form_data: dict = None):
        """
        Args:
            template_path: Chemin vers le template DOCX original
            form_data: Données du formulaire pour les placeholders
        """
        self.template_path = template_path
        self.doc = Document(template_path)
        self.form_data = form_data or {}
        
    def fill_with_ai(self, project_data: Dict[str, Any], evidence_pack: str, 
                     img_catalog: List[Dict], openai_client, model: str) -> Document:
        """
        Remplit le template en utilisant l'IA pour analyser les pièces.
        
        Args:
            project_data: Données du projet (formulaire)
            evidence_pack: Extraits des pièces uploadées
            img_catalog: Liste des images disponibles
            openai_client: Client OpenAI configuré
            model: Nom du modèle à utiliser
            
        Returns:
            Document DOCX rempli
        """
        # 0. Remplacer les placeholders du formulaire AVANT l'IA
        logger.info("[PPSPS] Étape 0: Remplacement des placeholders du formulaire")
        logger.info(f"[PPSPS] form_data reçu: {self.form_data}")
        placeholder_values = self._prepare_placeholder_values()
        replace_placeholders_in_doc(self.doc, placeholder_values)
        logger.info("[PPSPS] Placeholders remplacés, passage à l'IA")
        
        # 1. Créer le prompt pour l'IA
        prompt = self._build_fill_prompt(project_data, evidence_pack, img_catalog)
        
        # 2. Appeler l'IA pour obtenir les données de remplissage
        messages = [
            {
                "role": "system",
                "content": "Tu es un expert coordinateur SPS qui remplit des PPSPS. "
                          "Tu analyses les documents fournis et extrais les informations pertinentes. "
                          "Tu réponds UNIQUEMENT en JSON valide."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
        
        response = openai_client.chat.completions.create(
            model=model,
            temperature=0.2,
            messages=messages
        )
        
        # 3. Parser la réponse JSON
        raw_response = response.choices[0].message.content.strip()
        # Nettoyer les backticks markdown si présents
        raw_response = raw_response.replace('```json', '').replace('```', '').strip()
        fill_data = json.loads(raw_response)
        
        # 4. Remplir le template avec les données
        self._fill_document(fill_data, img_catalog)
        
        return self.doc
    
    def _build_fill_prompt(self, project_data: Dict, evidence_pack: str, 
                          img_catalog: List[Dict]) -> str:
        """Construit le prompt pour guider l'IA dans le remplissage."""
        
        return f"""Tu dois extraire et structurer les informations pour remplir un PPSPS.

🎯 RÈGLE ABSOLUE : PRIORITÉ DES SOURCES
1. **TOUJOURS utiliser EN PRIORITÉ les informations des PIÈCES UPLOADÉES** (extraits ci-dessous)
2. Le **formulaire** sert UNIQUEMENT de **FALLBACK** si l'info est absente des pièces
3. Si une info est trouvée dans les pièces, l'utiliser MÊME si le formulaire contient autre chose

📄 PIÈCES UPLOADÉES (PRIORITÉ ABSOLUE) :
{evidence_pack}

📝 FORMULAIRE (FALLBACK UNIQUEMENT) :
{json.dumps(project_data, ensure_ascii=False, indent=2)}

🖼️ IMAGES DISPONIBLES :
{json.dumps(img_catalog, ensure_ascii=False, indent=2)}

Réponds en JSON avec cette structure EXACTE :

{{
  "informations_generales": {{
    "nom_entreprise": "...",
    "telephone": "...",
    "adresse": "...",
    "email": "...",
    "fax": "...",
    "nom_chef_entreprise": "...",
    "description_operation": "Description détaillée de l'opération/chantier",
    "lot": "Lot de l'entreprise",
    "travaux_confies": "Description des travaux confiés à l'entreprise",
    "date_debut": "JJ/MM/AAAA",
    "date_fin": "JJ/MM/AAAA",
    "effectif_moyen": "nombre",
    "effectif_pointe": "nombre"
  }},
  
  "organismes_prevention": {{
    "medecine_travail": {{"nom": "...", "telephone": "..."}},
    "inspecteur_travail": {{"nom": "...", "telephone": "..."}},
    "csps": {{"nom": "...", "telephone": "..."}},
    "carsat": {{"nom": "...", "telephone": "..."}}
  }},
  
  "mesures_hygiene": {{
    "vestiaires": {{
      "description": "Description des vestiaires",
      "emplacement": "Lieu",
      "date_service": "JJ/MM/AAAA"
    }},
    "sanitaires": {{
      "description": "Description des sanitaires",
      "emplacement": "Lieu",
      "date_service": "JJ/MM/AAAA"
    }},
    "restauration": {{
      "description": "Description de la restauration",
      "emplacement": "Lieu",
      "date_service": "JJ/MM/AAAA"
    }}
  }},
  
  "secours_evacuation": {{
    "pompiers": "18 ou 112",
    "samu": "15",
    "police": "17",
    "centre_antipoison": "...",
    "sst_chantier": [{{"nom": "...", "telephone": "..."}}],
    "point_rassemblement": "Description du point de rassemblement",
    "consignes_specifiques": "Consignes spécifiques au chantier"
  }},
  
  "risques_travaux": [
    {{
      "phase": "Phase de travail 1",
      "moyens": "Matériels, équipements utilisés",
      "risques_entreprise": "Risques pour nos salariés",
      "risques_autres": "Risques pour les autres intervenants",
      "prevention": "Mesures de prévention mises en place"
    }}
  ],
  
  "risques_environnement": [
    {{
      "categorie": "Déplacements du personnel sur le chantier",
      "contraintes_environnement": "Contraintes liées à l'environnement",
      "risques_autres_intervenants": "Risques des autres intervenants",
      "prevention": "Moyens de prévention",
      "observations": "Observations éventuelles"
    }},
    {{
      "categorie": "Organisation du chantier",
      "contraintes_environnement": "...",
      "risques_autres_intervenants": "...",
      "prevention": "...",
      "observations": "..."
    }},
    {{
      "categorie": "Autres",
      "contraintes_environnement": "...",
      "risques_autres_intervenants": "...",
      "prevention": "...",
      "observations": "..."
    }}
  ],
  
  "annexes": [
    {{
      "titre": "Plans de circulation",
      "images": ["fichier1.png", "fichier2.png"],
      "description": "Description optionnelle"
    }}
  ]
}}

⚠️ RÈGLES IMPORTANTES :
- Si une info n'est pas trouvée : laisser chaîne vide "" ou null
- Pour les téléphones : format exact trouvé dans les docs (ex: "01 23 45 67 89")
- Pour les dates : format JJ/MM/AAAA
- Pour les risques : être FACTUEL et PRÉCIS (pas de généralités)
- Séparer les différents risques/phases avec des détails distincts
- Pour les annexes : utiliser UNIQUEMENT les noms de fichiers du catalogue fourni
- Si pas d'info dans les pièces ET dans le formulaire : laisser vide
"""
    
    def _fill_document(self, fill_data: Dict, img_catalog: List[Dict]):
        """Remplit le document avec les données extraites par l'IA."""
        
        # 1. Remplir le premier tableau (informations générales)
        if len(self.doc.tables) > 0:
            self._fill_info_table(self.doc.tables[0], fill_data.get("informations_generales", {}))
        
        # 2. Remplir le deuxième tableau (description opération)
        if len(self.doc.tables) > 1:
            self._fill_operation_table(self.doc.tables[1], fill_data.get("informations_generales", {}))
        
        # 3. Remplir le tableau hygiène
        if len(self.doc.tables) > 2:
            self._fill_hygiene_table(self.doc.tables[2], fill_data.get("mesures_hygiene", {}))
        
        # 4. Remplir le tableau risques travaux
        if len(self.doc.tables) > 3:
            self._fill_risques_travaux_table(self.doc.tables[3], 
                                            fill_data.get("risques_travaux", []))
        
        # 5. Remplir le tableau risques environnement
        if len(self.doc.tables) > 4:
            self._fill_risques_env_table(self.doc.tables[4], 
                                        fill_data.get("risques_environnement", []))
        
        # 6. Remplir les sections texte (organismes, secours)
        self._fill_text_sections(fill_data)
        
        # 7. Ajouter les annexes à la fin
        self._add_annexes(fill_data.get("annexes", []), img_catalog)
    
    
    def _prepare_placeholder_values(self) -> dict:
        """Prépare les valeurs des placeholders à partir des données du formulaire."""
        values = {
            "{{NOM_PROJET}}": self.form_data.get("name", ""),
            "{{ADRESSE_CHANTIER}}": self.form_data.get("address", ""),
            "{{REFERENCE_AFFAIRE}}": self.form_data.get("project_reference", ""),
            "{{TELEPHONE_CHANTIER}}": self.form_data.get("site_phone", ""),
            "{{DUREE_SEMAINES}}": str(self.form_data.get("duration_weeks", "")) if self.form_data.get("duration_weeks") else "",
            "{{EFFECTIF_MAXIMUM}}": str(self.form_data.get("workforce", "")) if self.form_data.get("workforce") else "",
            "{{LOTS_TRAVAUX}}": self.form_data.get("works_csv", ""),
            "{{NOM_ENTREPRISE}}": self.form_data.get("company_name", ""),
            "{{ADRESSE_ENTREPRISE}}": self.form_data.get("company_address", ""),
            "{{TELEPHONE_ENTREPRISE}}": self.form_data.get("company_phone", ""),
            "{{EMAIL_ENTREPRISE}}": self.form_data.get("company_email", ""),
            "{{RESPONSABLE_TRAVAUX}}": self.form_data.get("site_manager_name", ""),
            "{{MAITRE_OUVRAGE}}": self.form_data.get("owner_name", ""),
            "{{MAITRE_OEUVRE}}": self.form_data.get("architect_name", ""),
        }
        
        # LOG: Afficher les valeurs des placeholders
        logger.info("[PLACEHOLDERS] Valeurs préparées:")
        for key, val in values.items():
            logger.info(f"  {key} = '{val}'")
        
        return values

    def _fill_info_table(self, table, data: Dict):
        """Remplit le tableau d'informations générales (TABLE 0)."""
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
            cell = table.rows[0].cells[0]
            
            # Construire le texte avec les données
            nom = data.get("nom_entreprise", "")
            tel = data.get("telephone", "")
            adresse = data.get("adresse", "")
            email = data.get("email", "")
            fax = data.get("fax", "")
            chef = data.get("nom_chef_entreprise", "")
            
            new_text = (
                f"Nom de l'entreprise : {nom if nom else '…' * 40}\n"
                f"Tél. : {tel if tel else '…' * 20}\n"
                f"Adresse : {adresse if adresse else '…' * 50}\n"
                f"E-mail : {email if email else '…' * 30}\n"
                f"Fax : {fax if fax else '…' * 30}\n"
                f"Nom du Chef d'entreprise : {chef if chef else '…' * 40}"
            )
            
            cell.text = new_text
    
    def _fill_operation_table(self, table, data: Dict):
        """Remplit le tableau description de l'opération (TABLE 1)."""
        if len(table.rows) >= 4:
            # Ligne 0 : Description et Lot
            if len(table.rows[0].cells) >= 3:
                desc = data.get("description_operation", "")
                lot = data.get("lot", "")
                table.rows[0].cells[1].text = desc if desc else ""
                table.rows[0].cells[2].text = f"Lot : {lot if lot else ''}"
            
            # Ligne 1 : Travaux confiés
            if len(table.rows[1].cells) >= 2:
                travaux = data.get("travaux_confies", "")
                table.rows[1].cells[1].text = travaux if travaux else ""
            
            # Ligne 2 : Planning
            if len(table.rows[2].cells) >= 2:
                debut = data.get("date_debut", "")
                fin = data.get("date_fin", "")
                table.rows[2].cells[1].text = f"Date de début : {debut}\tDate de fin : {fin}"
            
            # Ligne 3 : Effectifs
            if len(table.rows[3].cells) >= 2:
                moyen = data.get("effectif_moyen", "")
                pointe = data.get("effectif_pointe", "")
                table.rows[3].cells[1].text = f"Effectif moyen : {moyen}\tEffectif de pointe : {pointe}"
    
    def _fill_hygiene_table(self, table, data: Dict):
        """Remplit le tableau mesures d'hygiène (TABLE 2)."""
        # Vestiaires (lignes 0-2)
        if len(table.rows) > 2:
            vest = data.get("vestiaires", {})
            if len(table.rows[1].cells) >= 2:
                table.rows[1].cells[1].text = vest.get("description", "")
            if len(table.rows[2].cells) >= 3:
                table.rows[2].cells[1].text = vest.get("emplacement", "")
                table.rows[2].cells[2].text = f"Date de mise en service : {vest.get('date_service', '')}"
        
        # Sanitaires (lignes 3-5)
        if len(table.rows) > 5:
            sani = data.get("sanitaires", {})
            if len(table.rows[4].cells) >= 2:
                table.rows[4].cells[1].text = sani.get("description", "")
            if len(table.rows[5].cells) >= 3:
                table.rows[5].cells[1].text = sani.get("emplacement", "")
                table.rows[5].cells[2].text = f"Date de mise en service : {sani.get('date_service', '')}"
        
        # Restauration (lignes 6-8)
        if len(table.rows) > 8:
            resto = data.get("restauration", {})
            if len(table.rows[7].cells) >= 2:
                table.rows[7].cells[1].text = resto.get("description", "")
            if len(table.rows[8].cells) >= 3:
                table.rows[8].cells[1].text = resto.get("emplacement", "")
                table.rows[8].cells[2].text = f"Date de mise en service : {resto.get('date_service', '')}"
    
    def _fill_risques_travaux_table(self, table, risques: List[Dict]):
        """Remplit le tableau d'analyse des risques travaux (TABLE 3)."""
        if not risques or len(table.rows) < 3:
            return
        
        # Supprimer les lignes existantes après l'en-tête (garder lignes 0-1)
        while len(table.rows) > 2:
            table._element.remove(table.rows[-1]._element)
        
        # Ajouter une ligne par phase/risque
        for risque in risques:
            row = table.add_row()
            cells = row.cells
            
            if len(cells) >= 5:
                cells[0].text = risque.get("phase", "")
                cells[1].text = risque.get("moyens", "")
                cells[2].text = risque.get("risques_entreprise", "")
                cells[3].text = risque.get("risques_autres", "")
                cells[4].text = risque.get("prevention", "")
                
                # Ajouter une bordure pointillée pour séparer les risques
                self._add_dotted_border(row)
    
    def _fill_risques_env_table(self, table, risques: List[Dict]):
        """Remplit le tableau risques liés à l'environnement (TABLE 4)."""
        if not risques or len(table.rows) < 2:
            return
        
        # Les 3 catégories fixes sont déjà dans le template (lignes 2-4)
        # On remplit juste les cellules correspondantes
        for i, risque in enumerate(risques[:3]):  # Max 3 catégories
            row_idx = 2 + i
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                if len(row.cells) >= 5:
                    # La première cellule contient déjà la catégorie
                    row.cells[1].text = risque.get("contraintes_environnement", "")
                    row.cells[2].text = risque.get("risques_autres_intervenants", "")
                    row.cells[3].text = risque.get("prevention", "")
                    row.cells[4].text = risque.get("observations", "")
    
    def _fill_text_sections(self, fill_data: Dict):
        """Remplit les sections textuelles (organismes, secours)."""
        orga = fill_data.get("organismes_prevention", {})
        secours = fill_data.get("secours_evacuation", {})
        
        # Parcourir les paragraphes et remplir aux bons endroits
        organismes_idx = None
        secours_idx = None
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip().upper()
            if "ORGANISMES DE PREVENTION" in text:
                organismes_idx = i
            elif "SECOURS ET EVACUATION" in text and organismes_idx is not None:
                secours_idx = i
            elif "MESURES" in text or "HYGIENE" in text:
                break  # On s'arrête après avoir trouvé les sections
        
        # Remplir ORGANISMES DE PREVENTION (paragraphes juste après le titre)
        if organismes_idx is not None and orga:
            insert_idx = organismes_idx + 1
            content = []
            
            med = orga.get("medecine_travail", {})
            if med.get("nom") or med.get("telephone"):
                content.append(f"Médecine du travail : {med.get('nom', '')} - Tél : {med.get('telephone', '')}")
            
            insp = orga.get("inspecteur_travail", {})
            if insp.get("nom") or insp.get("telephone"):
                content.append(f"Inspection du travail : {insp.get('nom', '')} - Tél : {insp.get('telephone', '')}")
            
            csps = orga.get("csps", {})
            if csps.get("nom") or csps.get("telephone"):
                content.append(f"CSPS : {csps.get('nom', '')} - Tél : {csps.get('telephone', '')}")
            
            carsat = orga.get("carsat", {})
            if carsat.get("nom") or carsat.get("telephone"):
                content.append(f"CARSAT : {carsat.get('nom', '')} - Tél : {carsat.get('telephone', '')}")
            
            if content:
                # Remplir ou créer le paragraphe
                if insert_idx < len(self.doc.paragraphs):
                    self.doc.paragraphs[insert_idx].text = "\n".join(content)
                else:
                    p = self.doc.add_paragraph("\n".join(content))
        
        # Remplir SECOURS ET EVACUATION (paragraphe juste après le titre)
        if secours_idx is not None and secours:
            insert_idx = secours_idx + 1
            content = []
            
            if secours.get("pompiers"):
                content.append(f"Pompiers : {secours.get('pompiers')}")
            if secours.get("samu"):
                content.append(f"SAMU : {secours.get('samu')}")
            if secours.get("police"):
                content.append(f"Police : {secours.get('police')}")
            if secours.get("centre_antipoison"):
                content.append(f"Centre antipoison : {secours.get('centre_antipoison')}")
            
            sst_list = secours.get("sst_chantier", [])
            if sst_list:
                content.append("\nSauveteurs Secouristes du Travail (SST) :")
                for sst in sst_list:
                    if isinstance(sst, dict):
                        content.append(f"  - {sst.get('nom', '')} : {sst.get('telephone', '')}")
            
            if secours.get("point_rassemblement"):
                content.append(f"\nPoint de rassemblement : {secours.get('point_rassemblement')}")
            
            if secours.get("consignes_specifiques"):
                content.append(f"\nConsignes spécifiques :\n{secours.get('consignes_specifiques')}")
            
            if content:
                # Remplir ou créer le paragraphe
                if insert_idx < len(self.doc.paragraphs):
                    self.doc.paragraphs[insert_idx].text = "\n".join(content)
                else:
                    p = self.doc.add_paragraph("\n".join(content))
    
    def _add_annexes(self, annexes: List[Dict], img_catalog: List[Dict]):
        """Ajoute les annexes à la fin du document."""
        
        # Trouver la section ANNEXES
        annexes_found = False
        for para in self.doc.paragraphs:
            if "ANNEXES" in para.text and any(run.bold for run in para.runs):
                annexes_found = True
                break
        
        if not annexes_found:
            # Ajouter le titre ANNEXES
            self.doc.add_page_break()
            p = self.doc.add_paragraph()
            run = p.add_run("ANNEXES")
            run.bold = True
            run.font.size = Pt(14)
        
        # Ajouter chaque annexe
        for annexe in annexes:
            titre = annexe.get("titre", "")
            images = annexe.get("images", [])
            description = annexe.get("description", "")
            
            if titre:
                p = self.doc.add_paragraph()
                run = p.add_run(f"\n{titre}")
                run.bold = True
                run.font.size = Pt(12)
            
            if description:
                self.doc.add_paragraph(description)
            
            # Ajouter les images
            for img_name in images:
                # Trouver le chemin complet de l'image
                img_path = None
                for img in img_catalog:
                    if img.get("file") == img_name:
                        img_path = img.get("stored_path")
                        break
                
                if img_path:
                    try:
                        from docx.shared import Inches
                        self.doc.add_picture(img_path, width=Inches(6))
                        # Légende
                        p = self.doc.add_paragraph(f"Figure : {img_name}")
                        p.alignment = 1  # Centré
                    except Exception:
                        # Si erreur, ajouter juste une mention
                        self.doc.add_paragraph(f"[Image : {img_name}]")
    
    def _add_dotted_border(self, row):
        """Ajoute une bordure pointillée en bas d'une ligne de tableau."""
        tcPr = row.cells[0]._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'dotted')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '808080')
        tcBorders.append(bottom)
        tcPr.append(tcBorders)


# ===== App =====
ENV = os.getenv("ENV", "dev").lower()  # "prod" en production
HTTPS_ONLY = (ENV == "prod")

app = FastAPI(
    title=APP_TITLE,
    docs_url=None,
    redoc_url=None,
    openapi_url=None
)

app.add_middleware(
    SessionMiddleware,
    secret_key=SECRET_KEY,
    same_site="lax",
    https_only=HTTPS_ONLY,   # True en prod, False en dev
    max_age=60 * 60 * 24 * 7,
)


FRONTEND_ORIGINS = [o.strip() for o in os.getenv("FRONTEND_ORIGINS", "").split(",") if o.strip()]
app.add_middleware(
    CORSMiddleware,
    allow_origins=FRONTEND_ORIGINS if FRONTEND_ORIGINS else [],  # vide => même origine
    allow_credentials=True,
    allow_methods=["GET","POST","DELETE","OPTIONS"],
    allow_headers=["Authorization","Content-Type","X-CSRF-Token"],
)


templates = Jinja2Templates(directory="templates")

# Mount static files (pour images, CSS, JS, etc.)
app.mount("/static", StaticFiles(directory="static"), name="static")

import secrets

def _ensure_csrf_token(request: Request) -> str:
    tok = request.session.get("csrf_token")
    if not tok:
        tok = secrets.token_hex(16)
        request.session["csrf_token"] = tok
    return tok

def _check_csrf(request: Request, token_from_form: str | None):
    expected = request.session.get("csrf_token")
    if not expected or not token_from_form or token_from_form != expected:
        raise HTTPException(status_code=400, detail="CSRF token invalide")


@app.on_event("startup")
def on_startup():
    init_db()
    # Créer le package de tokens par défaut
    with Session(engine) as session:
        StripeService.create_default_packages(session)

# =====================================================================
#                               AUTH
# =====================================================================

from fastapi.responses import RedirectResponse

pwd_ctx = CryptContext(
    schemes=["argon2"],
    deprecated="auto",
)

def hash_pwd(p: str) -> str:
    return pwd_ctx.hash(p)

def verify_pwd(p: str, h: str) -> bool:
    return pwd_ctx.verify(p, h)

def get_current_user(request: Request, session: Session = Depends(get_session)) -> Optional[UserDB]:
    uid = request.session.get("uid")
    if not uid:
        return None
    return session.get(UserDB, uid)

def require_login(request: Request, session: Session = Depends(get_session)) -> UserDB:
    uid = request.session.get("uid")
    if not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user = session.get(UserDB, uid)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return user

def _require_admin(request: Request, session: Session = Depends(get_session)):
    uid = request.session.get("uid")
    if not uid:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user = session.get(UserDB, uid)
    if not user or not getattr(user, "is_admin", False):
        raise HTTPException(status_code=403, detail="Forbidden")

# =====================================================================
#                              SCHEMAS API
# =====================================================================
class Project(BaseModel):
    name: str = PydField(min_length=2, max_length=200)
    address: str = PydField(min_length=3, max_length=300)
    works: list[str] = PydField(default_factory=list)
    duration_weeks: int = PydField(ge=0, le=520)
    workforce: int = PydField(ge=0, le=5000)
    companies: list[str] = PydField(default_factory=list)

    @field_validator("works", mode="before")
    @classmethod
    def normalize_works(cls, v):
        if not v: return []
        return [str(x).strip()[:80] for x in v][:20]

class DocumentIn(BaseModel):
    doc_type: Literal["PPSPS"]
    content_md: str

class ProjectFactsIn(BaseModel):
    site_env: str = ""
    zones: list[str] = PydField(default_factory=list)
    eu_name: str = ""
    csps_name: str = ""
    start_date: Optional[str] = None  # "YYYY-MM-DD"
    end_date: Optional[str] = None
    work_hours: str = ""

# =====================================================================
#                           UTILS (fichiers)
# =====================================================================
ALLOWED_EXTS = {".pdf", ".docx", ".txt", ".md"}
MAX_FILE_MB = 25

def _safe_name(name: str) -> str:
    return "".join(c for c in name if c.isalnum() or c in (".", "_", "-", " ")).strip()

def _project_upload_dir(project_id: int) -> str:
    # retourner juste "<id>" pour être combiné avec UPLOADS_ROOT
    d = os.path.join(str(project_id))
    os.makedirs(os.path.join("uploads", d), exist_ok=True)  # garde la création physique si tu veux
    return d

def _ext(path: str) -> str:
    return os.path.splitext(path)[1].lower()

def extract_text_from_file(path: str) -> str:
    ext = _ext(path)
    try:
        if ext == ".pdf":
            out = []
            with open(path, "rb") as f:
                pdf = PdfReader(f)
                for page in pdf.pages:
                    out.append(page.extract_text() or "")
            return "\n".join(out)
        elif ext == ".docx":
            doc = DocxDocument(path)
            return "\n".join(p.text for p in doc.paragraphs)
        elif ext in (".txt", ".md"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        else:
            return ""
    except Exception:
        return ""
    
    # --- Helpers sécurité & rate-limit ---

UPLOADS_ROOT = os.path.abspath("uploads")
_RATE_BUCKET = {}  # { key: [timestamps_sec...] }

def _safe_path(base: str, path: str) -> str:
    """Empêche path traversal : renvoie un chemin normalisé OBLIGATOIREMENT sous base."""
    base_abs = os.path.abspath(base)
    full = os.path.abspath(os.path.normpath(os.path.join(base_abs, path)))
    if not full.startswith(base_abs + os.sep) and full != base_abs:
        raise HTTPException(status_code=400, detail="Chemin invalide")
    return full

def _require_rate_limit(key: str, max_calls: int = 1, window_sec: int = 60):
    """Limite simple en mémoire (par worker) : max_calls / window_sec."""
    now = int(datetime.utcnow().timestamp())
    bucket = _RATE_BUCKET.get(key, [])
    bucket = [t for t in bucket if now - t < window_sec]
    if len(bucket) >= max_calls:
        raise HTTPException(status_code=429, detail="Trop de requêtes, réessaie bientôt.")
    bucket.append(now)
    _RATE_BUCKET[key] = bucket

def _internal_error(msg: str = "Erreur interne."):
    # Évite de renvoyer des détails sensibles (stack, messages upstream, etc.)
    raise HTTPException(status_code=500, detail=msg)


# =====================================================================
#        DÉTECTION DE LOTS (heuristique)
# =====================================================================
WORK_KEYWORDS = {
    "toiture": ["toiture", "étanchéité", "charpente", "couverture"],
    "levage": ["grue", "levage", "manutention lourde", "palonnier"],
    "électrique": ["électricité", "tableau", "câblage", "TGBT", "consignation"],
    "échafaudage": ["échafaudage", "plancher de travail"],
    "peinture": ["peinture", "revêtement"],
    "maçonnerie": ["maçonnerie", "béton", "coffrage"],
    "plomberie": ["plomberie", "PVC", "cuivre", "vanne", "collecteur"],
    "climatisation": ["CVC", "climatisation", "ventilation", "CTA", "gaines"],
    "soudure": ["soudure", "meulage", "travaux à chaud", "chalumeau"],
}

def _check_files_relevance_with_ai(session: Session, project_id: int) -> tuple[bool, str]:
    """
    Utilise l'IA pour déterminer si les fichiers sont pertinents pour un PPSPS.
    Plus intelligent et flexible qu'une simple détection de mots-clés.
    
    Returns:
        (is_relevant: bool, message: str)
        - Si is_relevant=False : message contient la raison du rejet
        - Si is_relevant=True et message non vide : avertissement (génération autorisée)
        - Si is_relevant=True et message vide : tout est OK
    """
    # 1. Récupérer le contenu des fichiers (tous concaténés)
    blob = _project_text_blob(session, project_id, limit_chars=15_000)
    
    if not blob or len(blob.strip()) < 200:
        return False, (
            "Les fichiers uploadés ne contiennent pas assez de texte exploitable. "
            "Vérifiez que vos documents sont lisibles et pertinents (PGC, plans, DICT, etc.)"
        )
    
    # 2. Construire le prompt pour l'IA
    prompt = f"""Tu es un expert en prévention BTP et coordinateur SPS. Analyse le contenu ci-dessous et détermine s'il est pertinent pour générer un PPSPS (Plan Particulier de Sécurité et de Protection de la Santé).

Documents PERTINENTS pour un PPSPS (accepter) :
- PGC (Plan Général de Coordination)
- PPSPS existants ou brouillons
- Plans de prévention
- Plans de circulation / PICH
- DICT/DT (déclarations de réseaux)
- Plans d'installations de chantier
- Documents techniques (VRD, gros œuvre, etc.)
- Fiches de Données de Sécurité (FDS)
- Notices de postes, modes opératoires
- Plans architecte, plans masse
- CCTP, cahier des charges travaux

Documents NON PERTINENTS (rejeter) :
- CV / lettres de motivation
- Factures / devis génériques sans lien technique
- Contrats commerciaux / RIB / statuts société
- Documents administratifs sans lien avec le chantier
- Manuels utilisateur sans rapport avec le BTP
- Documents personnels

CONTENU À ANALYSER :
{blob[:12000]}

Réponds UNIQUEMENT par un JSON au format exact suivant (pas de texte avant/après) :
{{
  "pertinent": true ou false,
  "confiance": nombre entre 0 et 100,
  "types_detectes": ["type1", "type2"],
  "raison": "Explication courte en 1-2 phrases maximum"
}}

Règles :
- Si les documents semblent liés au BTP/chantier : pertinent=true
- Si uniquement documents administratifs/RH : pertinent=false
- Si doute : pertinent=true mais confiance < 70
- types_detectes : liste précise (ex: ["PGC", "plan circulation"])
- raison : courte et claire"""

    # 3. Appel à l'IA
    try:
        response = client.chat.completions.create(
            model=OPENROUTER_DEFAULT_MODEL,
            temperature=0.1,  # Peu de créativité pour plus de précision
            max_tokens=300,
            messages=[
                {
                    "role": "system", 
                    "content": "Tu es un expert BTP. Réponds UNIQUEMENT en JSON valide, sans texte avant/après."
                },
                {"role": "user", "content": prompt}
            ]
        )
        
        raw = response.choices[0].message.content.strip()
        
        # Nettoyer les éventuels backticks markdown
        raw = raw.replace('```json', '').replace('```', '').strip()
        
        # Vérifier que la réponse n'est pas vide
        if not raw:
            logger.warning("[RELEVANCE] Réponse vide de l'IA, on laisse passer")
            return True, ""
        
        # Parser le JSON
        result = json.loads(raw)
        
        pertinent = result.get("pertinent", False)
        confiance = result.get("confiance", 0)
        types_detectes = result.get("types_detectes", [])
        raison = result.get("raison", "Raison inconnue")
        
        # Log pour debug
        logger.info(f"[RELEVANCE] Pertinent={pertinent}, Confiance={confiance}%, Types={types_detectes}")
        
        # 4. Décision finale
        if not pertinent:
            # REJET : Documents clairement non pertinents
            return False, (
                f"❌ Documents non pertinents pour un PPSPS.\n\n"
                f"Raison : {raison}\n\n"
                f"Veuillez uploader des documents liés au chantier : PGC, plans de prévention, "
                f"DICT, plans de circulation, documents techniques, etc."
            )
        
        # ACCEPTATION avec avertissement si confiance faible
        if confiance < 70:
            types_str = ", ".join(types_detectes) if types_detectes else "documents partiels"
            return True, (
                f"⚠️ Documents détectés ({types_str}) mais qualité incertaine (confiance {confiance}%). "
                f"Raison : {raison}\n"
                f"La génération est autorisée mais vérifiez le résultat attentivement."
            )
        
        # ACCEPTATION totale
        return True, ""
        
    except json.JSONDecodeError as e:
        # Si l'IA n'a pas renvoyé du JSON valide
        logger.error(f"[RELEVANCE] JSON invalide : {e}")
        # Principe de précaution : on accepte mais on log l'erreur
        return True, ""
        
    except Exception as e:
        # Erreur technique (API, timeout, etc.)
        logger.error(f"[RELEVANCE] Erreur IA : {e}")
        # Principe de précaution : on laisse passer
        return True, ""

def detect_works_from_text(text: str) -> list[str]:
    t = (text or "").lower()
    found = set()
    for work, kws in WORK_KEYWORDS.items():
        if any(k.lower() in t for k in kws):
            found.add(work)
    return sorted(found)

def _project_text_blob(session: Session, project_id: int, limit_chars: int = 80_000) -> str:
    atts = session.exec(select(AttachmentDB).where(AttachmentDB.project_id == project_id)).all()
    if not atts:
        return ""
    parts = []
    for a in atts:
        if a.extracted_text:
            parts.append(f"\n\n===== FICHIER: {a.filename} =====\n{a.extracted_text}")
    blob = "\n".join(parts)
    return blob[:limit_chars]

# =====================================================================
#                     EXPORTS DOCX/PDF (Markdown → DOCX)
# =====================================================================
@app.post("/export_docx_stream")
def export_docx_stream(markdown: str = Body(...), filename: str = Body("document.docx")):
    # 1) DOCX de base
    d = DocxDocument()
    _build_doc_styles(d)

    # 2) Remplissage à partir du markdown découpé (texte / tableaux)
    segments = _split_text_and_tables(markdown or "")
    for kind, *payload in segments:
        if kind == "text":
            text = payload[0]
            # NOTE: adapte le nom si ta fonction s'appelle sans underscore
            _maybe_add_images(d, text, base_dir="uploads")
        elif kind == "table":
            section, csv_text = payload
            _docx_add_csv_table(d, section, csv_text)

    # 3) Pied de page ROBUSTE (ASCII only)
    section = d.sections[-1]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = "PPSPS genere par SPS Copilot - " + datetime.now().strftime("%Y-%m-%d %H:%M")

    # 4) Retour en streaming
    bio = BytesIO()
    d.save(bio)
    bio.seek(0)
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        bio,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )

def _img_to_data_url(path: str) -> str:
    mime, _ = mimetypes.guess_type(path)
    if not mime: mime = "image/png"
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("ascii")
    return f"data:{mime};base64,{b64}"

IMG_RANKER_SYSTEM = (
    "Tu es un coordonnateur SPS expert. Tu reçois des images extraites de pièces (PGC, plans...). "
    "Objectif: décider si chaque image est RÉELLEMENT pertinente pour un PPSPS et où l'insérer. "
    "ATTENTION : REJETTE IMPITOYABLEMENT les templates, fonds de page, logos, cadres vides."
)

IMG_RANKER_USER_TEMPLATE = """\
Analyse les images et renvoie un JSON strict, liste d'objets:
[
  {{
    "filename": "<nom du fichier>",
    "keep": true|false,
    "category": "plan_circulation"|"plan_levage"|"plan_reseaux"|"securite_secours"|"chimique_fds"|"autres",
    "suggested_location": "corps"|"Annexe A"|"Annexe B"|"Annexe C"|"Annexe D",
    "confidence": 0..1,
    "caption": "légende courte (<=120c)",
    "reason": "raison brève (<=120c)"
  }}
]

**RÈGLES DE TRI ULTRA-STRICTES** :

🚫 **REJETER SYSTÉMATIQUEMENT (keep=false)** :
- Couvertures, pages de garde, sommaires
- Logos, filigranes, en-têtes/pieds de page
- Cadres vides, gabarits, templates de mise en page
- Fonds décoratifs sans contenu technique
- Pages avec uniquement du texte (pas de schéma/plan)
- Snapshots de pages entières sans zoom sur un élément précis
- Images floues, illisibles ou de trop faible résolution
- Doublons ou images quasi-identiques

✅ **GARDER UNIQUEMENT (keep=true)** si **TOUS** ces critères :
1. Contenu technique EXPLOITABLE (plan, schéma, pictogramme, diagramme)
2. Lisibilité EXCELLENTE (texte/légendes lisibles, contraste suffisant)
3. Pertinence DIRECTE pour la sécurité/prévention du chantier
4. Pas de doublon avec une image déjà gardée
5. Confidence ≥ 0.80 (sinon rejeter)

**LIMITES STRICTES** :
- Maximum 2 images par catégorie (choisir les 2 meilleures)
- Si doute sur l'utilité : REJETER (principe de précaution)

**Catégories** :
- plan_circulation (PICH) → Annexe A + insertion section "Circulation"
- plan_levage → Annexe B + insertion section "Levage"
- plan_reseaux (DICT/DT) → Annexe C (mention seulement dans corps)
- securite_secours (évacuation/DAE) → insertion section "Secours"
- chimique_fds (CLP/FDS) → Annexe D (liste produits + renvoi)
- autres : rejeter sauf si indispensable (confidence ≥ 0.90)

Contexte projet : {context}

**RAPPEL** : Sois IMPITOYABLE. Mieux vaut 0 image que des templates inutiles.
"""

PEEK_SYSTEM = (
    "Tu es un coordonnateur SPS. Tu vois des miniatures de pages d'un PGC/PPSPS. "
    "But: ne garder que les pages qui contiennent des plans utiles au PPSPS."
)
PEEK_USER_INSTRUCTIONS = ("""\
Objectif: RENVOIE STRICTEMENT {"keep_pages":[...]} (1-based).
- Garde UNIQUEMENT : plans circulation/PICH, levage/manutention, réseaux (DICT/DT), évacuation/DAE, FDS/CLP lisibles.
- EXCLURE : couvertures, sommaires, tableaux administratifs sans schémas, logos/filigranes, gabarits/fonds, pages vides ou quasi textuelles.
- Ne renvoie que les n° de pages réellement utiles.
Exemple: {"keep_pages":[1,5,7]}
Contexte projet: {context}
""")



def _model_supports_vision() -> bool:
    m = (OPENROUTER_DEFAULT_MODEL or "").lower()
    return any(k in m for k in (
        "vision","vl","gpt-4o","qwen-vl","llava","gemini","claude-3.7",
        "gpt-5", "chatgpt-5", "gpt-5o"  # <- ajout
    )) or os.getenv("FORCE_VISION") == "1"


def peek_pages_for_plans(pdf_path: str, thumb_max_px: int = 384, batch: int = 12) -> list[int]:
    """Miniatures PNG base64 → LLM vision → pages à garder (1-based)."""
    if not _model_supports_vision():
        try:
            doc = fitz.open(pdf_path)
            keep = list(range(1, len(doc)+1))
            doc.close()
            return keep
        except Exception:
            return []

    thumbs = []
    doc = fitz.open(pdf_path)
    scale = thumb_max_px / 100.0
    for pno in range(len(doc)):
        pix = doc[pno].get_pixmap(matrix=fitz.Matrix(scale, scale))
        b64 = base64.b64encode(pix.tobytes("png")).decode("ascii")
        thumbs.append((pno + 1, f"data:image/png;base64,{b64}"))
    doc.close()

    keep_pages = set()
    for i in range(0, len(thumbs), batch):
        chunk = thumbs[i:i+batch]
        content = [{"type": "text", "text": PEEK_USER_INSTRUCTIONS}]
        for pno, url in chunk:
            content.append({"type": "text", "text": f"page: {pno}"})
            content.append({"type": "input_image", "image_url": url})
        try:
            resp = client.chat.completions.create(
                model=OPENROUTER_DEFAULT_MODEL, temperature=0,
                messages=[{"role":"system","content":PEEK_SYSTEM},
                          {"role":"user","content":content}]
            )
            raw = (resp.choices[0].message.content or "").strip()
            data = json.loads(raw)
            for n in (data.get("keep_pages") or []):
                if isinstance(n, int) and n >= 1:
                    keep_pages.add(n)
        except Exception:
            continue

    if not keep_pages:
        try:
            doc = fitz.open(pdf_path)
            keep = list(range(1, len(doc)+1))
            doc.close()
            return keep
        except Exception:
            return []
    return sorted(keep_pages)


def _rank_images_with_vision(session: Session, project_id: int, max_per_batch: int = 6) -> list[dict]:
    if not _model_supports_vision():
      return []
    
    # récup images (via DB)
    imgs = session.exec(
        select(AttachmentDB).where(
            AttachmentDB.project_id == project_id,
            AttachmentDB.mime_type.like("image%")
        ).order_by(AttachmentDB.created_at.asc())
    ).all()
    # pré-filtre (taille/dim)
    candidates = []
    for r in imgs:
        try:
            if os.path.getsize(r.stored_path) < 50_000:  # >50KB
                continue
            
            # Vérifier dimensions minimales
            from PIL import Image
            with Image.open(r.stored_path) as im:
                w, h = im.size
                # Au moins 400x400 pixels
                if w < 400 or h < 400:
                    continue
                # Ratio trop bizarre = probable template
                ratio = max(w, h) / min(w, h)
                if ratio > 4:  # Trop allongé
                    continue
        except Exception:
            continue
        candidates.append(r)

    # contexte léger
    proj = session.get(ProjectDB, project_id)
    ctx = {
        "name": proj.name, "address": proj.address,
        "works": (proj.works_csv or "").split(",")
    }

    results = []
    # batching pour éviter payload monstre
    for i in range(0, len(candidates), max_per_batch):
        batch = candidates[i:i+max_per_batch]
        content = [{"type": "text",
                    "text": IMG_RANKER_USER_TEMPLATE.format(context=json.dumps(ctx, ensure_ascii=False))}]
        for r in batch:
            content.append({"type": "input_image",
                            "image_url": _img_to_data_url(r.stored_path)})
            content.append({"type": "text", "text": f"filename: {r.filename}"})

        resp = client.chat.completions.create(
            model=OPENROUTER_DEFAULT_MODEL,  # ⚠️ doit être un modèle vision
            temperature=0,
            messages=[
                {"role": "system", "content": IMG_RANKER_SYSTEM},
                {"role": "user", "content": content}
            ]
        )
        raw = (resp.choices[0].message.content or "").strip()
        try:
            results.extend(json.loads(raw))
        except Exception:
            # si pas JSON, on ignore ce batch
            continue

    # borne par catégorie (évite le spam)
    caps = {"plan_circulation": 2, "plan_levage": 2, "plan_reseaux": 3, "securite_secours": 2, "chimique_fds": 3, "autres": 1}
    counts = {k:0 for k in caps}
    shortlisted = []
    for r in results:
        if not r.get("keep"):
            continue
        cat = r.get("category") or "autres"
        if counts.get(cat, 0) >= caps.get(cat, 0):
            continue
        counts[cat] += 1
        shortlisted.append(r)

    return shortlisted


# =====================================================================
#                          CRUD PROJETS / DOCS
# =====================================================================
@app.post("/projects/{project_id}/files/suggest-works")
def suggest_works_from_files(project_id: int, session: Session = Depends(get_session)):
    proj = session.get(ProjectDB, project_id)
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")
    atts = session.exec(select(AttachmentDB).where(AttachmentDB.project_id == project_id)).all()
    if not atts:
        return {"suggested": []}
    blob = "\n\n".join([a.extracted_text or "" for a in atts if a.extracted_text])
    found = detect_works_from_text(blob) if blob else []
    return {"suggested": found}

@app.post("/projects")
def create_project(p: Project, session: Session = Depends(get_session), user: UserDB = Depends(require_login)):
    project = ProjectDB(
        name=p.name,
        address=p.address,
        works_csv=",".join(p.works),
        duration_weeks=p.duration_weeks,
        workforce=p.workforce,
        companies_csv=",".join(p.companies),
        owner_id=user.id,
    )
    session.add(project)
    session.commit()
    session.refresh(project)
    return project

@app.get("/projects")
def list_projects(session: Session = Depends(get_session), user: UserDB = Depends(require_login)):
    return session.exec(select(ProjectDB).where(ProjectDB.owner_id == user.id)).all()

@app.post("/projects/{project_id}/documents")
def add_document(project_id: int, doc: DocumentIn, session: Session = Depends(get_session), user: UserDB = Depends(require_login)):
    proj = get_owned_project_or_404(project_id, user, session)
    existing = session.exec(
        select(DocumentDB).where(DocumentDB.project_id == proj.id, DocumentDB.doc_type == doc.doc_type)
    ).all()
    next_version = (max([d.version for d in existing]) + 1) if existing else 1
    dbdoc = DocumentDB(
        project_id=proj.id,
        doc_type=doc.doc_type,
        content_md=doc.content_md,
        version=next_version,
        status="draft",
    )
    session.add(dbdoc)
    session.commit()
    session.refresh(dbdoc)
    return dbdoc

@app.get("/projects/{project_id}/documents")
def list_documents(project_id: int, session: Session = Depends(get_session), user: UserDB = Depends(require_login)):
    proj = get_owned_project_or_404(project_id, user, session)
    return session.exec(select(DocumentDB).where(DocumentDB.project_id == proj.id)).all()

@app.get("/documents/{doc_id}")
def get_document(doc_id: int, session: Session = Depends(get_session)):
    doc = session.get(DocumentDB, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


def _add_cover_page(doc, project_name: str, project_address: str):
    """
    Ajoute une belle page de garde au début du document.
    """
    # Paragraphe vide pour espacer
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TITRE PRINCIPAL
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PPSPS")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(37, 99, 235)  # Bleu
    
    # SOUS-TITRE
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Plan Particulier de Sécurité et de Protection de la Santé")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(100, 116, 139)  # Gris
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # PROJET
    project_para = doc.add_paragraph()
    project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project_para.add_run(f"Projet : {project_name}")
    run.font.size = Pt(20)
    run.font.bold = True
    
    # ADRESSE
    addr_para = doc.add_paragraph()
    addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = addr_para.add_run(project_address)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # DATE
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Document généré le {datetime.now().strftime('%d/%m/%Y')}")
    run.font.size = Pt(11)
    run.font.italic = True
    
    # SAUT DE PAGE
    doc.add_page_break()


@app.get("/documents/{doc_id}/export_docx")
def export_docx_by_id(doc_id: int, session: Session = Depends(get_session), user: UserDB = Depends(require_login)):
    """
    Exporte le DOCX généré depuis le template.
    Récupère le fichier DOCX sauvegardé lors de la génération.
    """
    doc = ensure_doc_is_owned(doc_id, user, session)
    proj = session.get(ProjectDB, doc.project_id)
    
    # Extraire le chemin du DOCX depuis content_md
    # Format : "[DOCX généré : /path/to/file.docx]"
    import re
    match = re.search(r'\[DOCX généré : (.+?)\]', doc.content_md or '')
    
    if match:
        # Nouveau système : DOCX déjà généré
        docx_path = match.group(1)
        
        if os.path.exists(docx_path):
            # Retourner le DOCX sauvegardé
            filename = f"PPSPS_{proj.name.replace(' ', '_')}_{doc.version}.docx"
            headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
            
            with open(docx_path, 'rb') as f:
                bio = BytesIO(f.read())
            
            bio.seek(0)
            return StreamingResponse(
                bio,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers=headers
            )
        else:
            raise HTTPException(status_code=404, detail="DOCX généré introuvable sur le serveur")
    
    # Ancien système (fallback) : reconstruire depuis Markdown
    else:
        d = DocxDocument()
        _build_doc_styles(d)
        
        # Page de garde
        _add_cover_page(d, proj.name, proj.address)
        
        segments = _split_text_and_tables(doc.content_md or "")
        project_img_dir = os.path.join("uploads", str(doc.project_id), "images")
        img_lookup = _image_lookup_for_project(session, doc.project_id)
        
        for kind, *payload in segments:
            if kind == "text":
                text = payload[0]
                _maybe_add_images(d, text, base_dir=project_img_dir, image_lookup=img_lookup)
            elif kind == "table":
                section, csv_text = payload
                _docx_add_csv_table(d, section, csv_text)
        
        section = d.sections[-1]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = f"PPSPS - {proj.name} - Généré le {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        
        bio = BytesIO()
        d.save(bio)
        bio.seek(0)
        
        filename = f"PPSPS_{proj.name.replace(' ', '_')}_{doc.version}.docx"
        headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
        
        return StreamingResponse(
            bio,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )



import fitz  # PyMuPDF
from PIL import Image

def _images_outdir(project_id: int) -> str:
    d = os.path.join("uploads", str(project_id), "images")
    os.makedirs(d, exist_ok=True)
    return d

def _extract_images_from_docx(docx_path: str, out_dir: str) -> list[str]:
    """Extrait les images d'un DOCX en fichiers PNG."""
    try:
        zname = docx_path
        import zipfile, io
        out = []
        with zipfile.ZipFile(zname) as z:
            for name in z.namelist():
                if name.startswith("word/media/"):
                    data = z.read(name)
                    ext = os.path.splitext(name)[1].lower()
                    base = os.path.splitext(os.path.basename(zname))[0]
                    fn = f"IMG_{base}_{len(out)+1}.png"
                    fp = os.path.join(out_dir, fn)
                    try:
                        im = Image.open(io.BytesIO(data))
                        im.convert("RGB").save(fp, "PNG")
                        out.append(fp)
                    except Exception:
                        # si déjà PNG/JPG lisible, on écrit tel quel
                        with open(fp, "wb") as f: f.write(data)
                        out.append(fp)
        return out
    except Exception:
        return []

def _extract_images_from_pdf(pdf_path: str, out_dir: str, keep_pages: list[int] | None = None) -> list[tuple[str,int]]:
    """Extrait les images intégrées pour les pages autorisées. Snapshot HD seulement si aucune image intégrée."""
    out = []
    try:
        doc = fitz.open(pdf_path)
        base = os.path.splitext(os.path.basename(pdf_path))[0]
        pages_whitelist = set(keep_pages or range(1, len(doc)+1))
        for pno in range(len(doc)):
            page_no_1b = pno + 1
            if page_no_1b not in pages_whitelist:
                continue
            page = doc[pno]
            found = 0
            for i, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                if pix.n > 4:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                fp = os.path.join(out_dir, f"IMG_{base}_{page_no_1b}_{i+1}.png")
                pix.save(fp)
                out.append((fp, page_no_1b))
                found += 1
        doc.close()
    except Exception:
        pass
    return out


def _ensure_annexes_with_images(md: str, img_catalog: list[dict]) -> str:
    # si des [IMAGE:...] existent déjà, on ne touche pas
    if IMAGE_BLOCK_RE.search(md or ""):
        return md
    if not img_catalog:
        return md

    # groupe simple par tags → Annexe A/B/C/D/E
    groups = {"A": [], "B": [], "C": [], "D": [], "E": []}
    for it in img_catalog:
        f = it.get("file") or ""
        tags = it.get("tags") or []
        if "plan_circulation" in tags: groups["A"].append(f)
        elif "plan_levage" in tags: groups["B"].append(f)
        elif "plan_reseaux" in tags: groups["C"].append(f)
        else:
            # met le reste en D par défaut (FDS / divers)
            groups["D"].append(f)

    annexes_lines = ["\n\n# Annexes\n"]
    order = [("A", "Plan de circulation / PICH"),
             ("B", "Plan de levage / manutention"),
             ("C", "Plans réseaux / DICT-DT"),
             ("D", "FDS / documents divers"),
             ("E", "VGP / échafaudages")]
    for key, title in order:
        imgs = groups[key]
        annexes_lines.append(f"## Annexe {key} — {title}")
        if imgs:
            for f in imgs:
                annexes_lines.append(f"[IMAGE:{f}]")
        else:
            annexes_lines.append("Aucun document fourni")

    return (md or "") + "\n" + "\n".join(annexes_lines) + "\n"


# =====================================================================
#                             INGEST / FACTS LÉGERS
# =====================================================================
@app.post("/projects/{project_id}/facts")
def upsert_project_facts(project_id: int, data: ProjectFactsIn, session: Session = Depends(get_session)):
    proj = session.get(ProjectDB, project_id)
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")

    def _parse(d: Optional[str]):
        try:
            return date.fromisoformat(d) if d else None
        except Exception:
            return None

    if hasattr(proj, "site_env"): proj.site_env = (data.site_env or "").strip()
    if hasattr(proj, "zones_csv"): proj.zones_csv = ",".join([z.strip() for z in data.zones if z.strip()])
    if hasattr(proj, "eu_name"): proj.eu_name = (data.eu_name or "").strip()
    if hasattr(proj, "csps_name"): proj.csps_name = (data.csps_name or "").strip()
    if hasattr(proj, "work_hours"): proj.work_hours = (data.work_hours or "").strip()
    if hasattr(proj, "start_date"): proj.start_date = _parse(data.start_date)
    if hasattr(proj, "end_date"): proj.end_date = _parse(data.end_date)

    proj.updated_at = datetime.utcnow()
    session.add(proj)
    session.commit()
    session.refresh(proj)
    return {"ok": True, "project_id": proj.id}

@app.post("/projects/{project_id}/ingest")
def ingest_project_files(project_id: int, session: Session = Depends(get_session)):
    """
    Ingestion = stockage + extraction texte + enrichissement 'works' (heuristique).
    Pas d'extraction JSON ni de binding template.
    """
    proj = session.get(ProjectDB, project_id)
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")

    blob = _project_text_blob(session, project_id)
    auto_from_text = detect_works_from_text(blob) if blob else []
    works = {w.strip() for w in (proj.works_csv or "").split(",") if w.strip()}
    works |= set(auto_from_text)
    proj.works_csv = ",".join(sorted(works))

    proj.updated_at = datetime.utcnow()
    session.add(proj)
    session.commit()
    session.refresh(proj)
    return {"ok": True, "project_id": proj.id, "works_csv": proj.works_csv}

def _docx_add_heading_or_paragraph(doc, line: str):
    s = line.strip()
    if not s:  # Skip les lignes vides
        return
    if s.startswith("# "):
        txt = s[2:].strip()
        # saut de page si H1 numéroté >= 2
        if re.match(r"^\d\.", txt) and not txt.startswith("1."):
            doc.add_page_break()
        p = doc.add_heading(txt, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif s.startswith("## "):
        doc.add_heading(s[3:].strip(), level=2)
    elif s.startswith("### "):
        doc.add_heading(s[4:].strip(), level=3)
    elif s == "---":  # Skip les séparateurs markdown
        return
    else:
        doc.add_paragraph(line)


# =====================================================================
#                KB Modes opératoires — Fallback texte
# =====================================================================
def _load_kb_modes_texts(base_dir: str = "app/kb/mesures") -> dict[str, str]:
    out = {}
    try:
        base = Path(__file__).resolve().parent / Path(base_dir)
        for fp in base.glob("*.md"):
            out[fp.stem.lower()] = fp.read_text(encoding="utf-8", errors="ignore").strip()
    except Exception:
        pass
    return out

KB_MAP = {
    # matériaux / substances
    "amiante": "amiante_plomb",
    "plomb": "amiante_plomb",
    "fibrociment": "amiante_plomb",

    # nuisances
    "bruit": "bruit",
    "sonore": "bruit",

    # travaux à chaud
    "chaud": "chaud",
    "soud": "chaud",
    "meul": "chaud",
    "chalumeau": "chaud",
    "flamme": "chaud",

    # chimique
    "chim": "chimique",
    "solvant": "chimique",
    "acide": "chimique",
    "résine": "chimique",
    "colle": "chimique",

    # circulation (interne / externe)
    "circulation": "circulation",
    "voirie": "circulation_externe",
    "route": "circulation_externe",
    "public": "circulation_externe",
    "livraison": "circulation_externe",

    # espaces confinés
    "confin": "confinement",
    "cuve": "confinement",
    "regard": "confinement",
    "réseau": "confinement",

    # électricité
    "élec": "electrique",
    "elec": "electrique",
    "tgbt": "electrique",
    "tableau": "electrique",
    "consign": "electrique",
    "vat": "electrique",
    "câbl": "electrique",

    # environnement
    "environ": "environnement",
    "déchet": "environnement",
    "pollution": "environnement",
    "rejet": "environnement",

    # hauteur / travail en hauteur
    "hauteur": "hauteur",
    "toiture": "hauteur",
    "étanch": "hauteur",
    "etanch": "hauteur",
    "chute": "hauteur",
    "acrotère": "hauteur",
    "echaf": "hauteur",

    # hygiène/base-vie
    "hyg": "hygiene",
    "sanit": "hygiene",
    "vestiaire": "hygiene",
    "réfect": "hygiene",

    # incendie/explosion
    "incend": "incendie_explosion",
    "feu": "incendie_explosion",
    "explos": "incendie_explosion",

    # levage
    "levage": "levage",
    "grue": "levage",
    "palonnier": "levage",
    "éling": "levage",

    # manutention
    "manut": "manutention",
    "charge": "manutention",
    "portage": "manutention",
    "transpalette": "manutention",

    # secours
    "secours": "secours",
    "urgence": "secours",
    "sst": "secours",
}

def _render_modes_ops_from_kb(project: ProjectDB) -> str:
    """
    Construit un bloc 'Modes opératoires' à partir de la KB si on n'a rien détecté dans les pièces.
    """
    kb_texts = _load_kb_modes_texts()
    if not kb_texts:
        return ""

    def _norm(s: str) -> str:
        s = (s or "").lower()
        rep = str.maketrans("áàâäãéèêëíìîïóòôöõúùûüç", "aaaaaeeeeiiiiooooouuuuc")
        return s.translate(rep)

    tokens = set()
    for w in (project.works_csv or "").split(","):
        w = _norm(w.strip())
        if w:
            tokens.add(w)

    chosen_keys = []
    for text in tokens:
        for needle, kb_key in KB_MAP.items():
            if needle in text:
                if kb_key in kb_texts and kb_key not in chosen_keys:
                    chosen_keys.append(kb_key)

    if not chosen_keys:
        return ""

    parts = ["**3. Modes opératoires**", ""]
    for key in chosen_keys:
        title = key.replace("_", " ").capitalize()
        body = kb_texts.get(key, "").strip()
        if not body:
            continue
        parts.append(f"### {title}\n{body}")

    return "\n\n".join(parts).strip()


def _build_doc_styles(doc):
    # police/parag défaut - Aptos moderne et lisible
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)

    # H1
    h1 = styles["Heading 1"]
    h1.font.name = "Aptos"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(31, 73, 125)  # Bleu foncé pro
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    # H2
    h2 = styles["Heading 2"]
    h2.font.name = "Aptos"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(31, 73, 125)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)

    # H3
    h3 = styles["Heading 3"]
    h3.font.name = "Aptos"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(3)

def _row_cant_split(row):
    # <w:cantSplit/> sur la ligne => évite coupure sur page
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement('w:cantSplit')
    trPr.append(cant)

def _format_table_pretty(tbl, header_fill=True):
    """Version améliorée avec couleurs."""
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # En-tête avec fond bleu
    if len(tbl.rows) > 0:
        for j, cell in enumerate(tbl.rows[0].cells):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)  # Texte blanc
            if header_fill:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), '2563EB')  # Bleu
                tcPr.append(shd)
    
    # Lignes alternées (gris très clair)
    for i, row in enumerate(tbl.rows[1:], start=1):
        if i % 2 == 0:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), 'F8FAFC')  # Gris très clair
                tcPr.append(shd)
    
    # Anti-coupure lignes
    for r in tbl.rows:
        _row_cant_split(r)


def _maybe_add_images(doc, text_block: str, base_dir: str | None = None, image_lookup: dict[str, str] | None = None):
    if not text_block:
        return
    pos = 0
    for m in IMAGE_BLOCK_RE.finditer(text_block):
        before = text_block[pos:m.start()]
        for line in before.splitlines():
            _docx_add_heading_or_paragraph(doc, line)

        img_ref = (m.group(1) or "").strip()
        img_path = img_ref

        # 1) lookup DB prioritaire (nom exact)
        if image_lookup and img_ref in image_lookup:
            img_path = image_lookup[img_ref]

        # 2) base_dir en secours
        elif base_dir and not os.path.isabs(img_path):
            cand = os.path.join(base_dir, img_ref)
            if os.path.exists(cand):
                img_path = cand

        if os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(6.5))
            except Exception as e:
                logger.warning(f"[IMG] Impossible de lire l'image {img_ref}: {e}")
        else:
            logger.warning(f"[IMG] Image introuvable, skippée: {img_ref}")

        pos = m.end()

    after = text_block[pos:]
    for line in after.splitlines():
        _docx_add_heading_or_paragraph(doc, line)


def _build_image_catalog(session: Session, project_id: int) -> list[dict]:
    rows = session.exec(
        select(AttachmentDB)
        .where(AttachmentDB.project_id == project_id)
        .order_by(AttachmentDB.created_at.asc())
    ).all()
    out = []
    for r in rows:
        if (r.mime_type or "").startswith("image"):
            # heuristique : tags depuis le nom
            name_l = (r.filename or "").lower()
            tag = []
            if "circul" in name_l or "pich" in name_l: tag.append("plan_circulation")
            if "levage" in name_l or "grue" in name_l: tag.append("plan_levage")
            if "dict" in name_l or "reseau" in name_l or "réseau" in name_l: tag.append("plan_reseaux")
            out.append({
                "file": r.filename,
                "stored_path": r.stored_path,
                "tags": tag,
                "note": (r.extracted_text or "")[:200]
            })
    return out


# =====================================================================
#                     EVIDENCE PACK & PROMPT FREEFORM
# =====================================================================

_MO_HINT_WORDS = (
    "mode opératoire", "procédure", "étapes", "balisage", "consignation",
    "échafaudage", "levage", "EPI", "sécurité", "travaux à chaud", "électricité"
)

def _scan_mo_candidates(blob: str, window_words: int = 220) -> list[str]:
    """
    Récupère des extraits candidats aux MOs autour de mots-clés.
    """
    if not blob:
        return []
    words = blob.split()
    text_lower = blob.lower()
    hits = []
    for hint in _MO_HINT_WORDS:
        idx = text_lower.find(hint.lower())
        if idx != -1:
            # approx index to word position
            # simple heuristic: split and take a window
            char_pos = idx
            # estimate word index by counting spaces up to char_pos
            before = blob[:char_pos]
            widx = before.count(" ")
            start = max(0, widx - window_words // 2)
            end = min(len(words), widx + window_words // 2)
            excerpt = " ".join(words[start:end]).strip()
            if excerpt and excerpt not in hits:
                hits.append(excerpt)
    return hits[:6]

TRAME_TITLES = [
    "1. Informations générales",
    "1.1 Affaire",
    "1.2 Chantier",
    "1.3 Acteurs & coordonnées",
    "1.4 Planning / durée / dates",
    "1.5 Effectifs / horaires",
    "1.6 Intervenants / Sous-traitants / matériel",
    "1.7 Installations, hygiène & conditions de travail",
    "2. Organisation du chantier",
    "2.1 Accès, circulation, balisage",
    "2.2 Installations techniques / base vie / sanitaires",
    "2.3 Conditions de travail & hygiène",
    "3. Modes opératoires",
    "4. Prévention / EPI / Risques",
    "5. Secours & évacuation",
    "6. Mise à jour / révisions du PPSPS",
    "7. Diffusion, consultation & conservation",
]

def _build_meta_hint(project: ProjectDB) -> dict:
    return {
        "project_name": project.name,
        "address": project.address,
        "works": [w.strip() for w in (project.works_csv or "").split(",") if w.strip()],
        "duration_weeks": project.duration_weeks,
        "workforce": project.workforce,
        "companies": [c.strip() for c in (project.companies_csv or "").split(",") if c.strip()],
        "site_env": getattr(project, "site_env", ""),
        "zones": [z.strip() for z in (getattr(project, "zones_csv", "") or "").split(",") if z.strip()],
        "eu_name": getattr(project, "eu_name", ""),
        "csps_name": getattr(project, "csps_name", ""),
        "work_hours": getattr(project, "work_hours", ""),
        "start_date": project.start_date.isoformat() if getattr(project, "start_date", None) else None,
        "end_date": project.end_date.isoformat() if getattr(project, "end_date", None) else None,
        # Nouveaux champs du formulaire
        "project_reference": getattr(project, "project_reference", ""),
        "site_phone": getattr(project, "site_phone", ""),
        "company_name": getattr(project, "company_name", ""),
        "company_address": getattr(project, "company_address", ""),
        "company_phone": getattr(project, "company_phone", ""),
        "company_email": getattr(project, "company_email", ""),
        "site_manager_name": getattr(project, "site_manager_name", ""),
        "owner_name": getattr(project, "owner_name", ""),
        "architect_name": getattr(project, "architect_name", ""),
    }

def _build_evidence_pack(blob: str, project: ProjectDB) -> str:
    """
    Structure l'evidence par grandes rubriques pour guider l'IA.
    """
    if not blob:
        blob = ""
    mos = _scan_mo_candidates(blob)
    lines = []
    lines.append("## Extraits — Informations générales\n")
    lines.append(blob[:6000])
    lines.append("\n\n## Extraits — Planning / Durée / Dates\n")
    lines.append(blob[6000:12000])
    lines.append("\n\n## Extraits — Effectifs / Sous-traitants / Matériel\n")
    lines.append(blob[12000:18000])
    lines.append("\n\n## Extraits — Conditions de travail / Hygiène / Installations\n")
    lines.append(blob[18000:24000])
    if mos:
        lines.append("\n\n## Extraits — Candidats Modes opératoires\n")
        for m in mos:
            lines.append(f"- {m}\n")
    return "\n".join(lines).strip()

def _prompt_freeform_ppsps(meta_hint: dict, evidence_pack: str, img_catalog: list[dict], kb_fallback_present: bool = True) -> list[dict]:

    """
    Construit les messages pour l'appel modèle (chat.completions).
    """
    TRAME_MD = """\
# 1. Informations générales
## 1.1 Affaire
## 1.2 Chantier
## 1.3 Acteurs & coordonnées
## 1.4 Planning / durée / dates
## 1.5 Effectifs / horaires
## 1.6 Intervenants / Sous-traitants / matériel
## 1.7 Installations, hygiène & conditions de travail

# 2. Organisation du chantier
## 2.1 Accès, circulation, balisage
## 2.2 Installations techniques / base vie / sanitaires
## 2.3 Conditions de travail & hygiène

# 3. Modes opératoires

# 4. Prévention / EPI / Risques

# 5. Secours & évacuation

# 6. Mise à jour / révisions du PPSPS

# 7. Diffusion, consultation & conservation
"""

    RULES = f"""
Tu es **rédacteur PPSPS expert**.
Objectif : produire un **PPSPS complet en français**, **au format Markdown**, qui respecte **strictement** la trame ci-dessous.

🚨 **RÈGLE ABSOLUE N°1 — PRIORITÉ DES SOURCES** 🚨
1. **TOUJOURS utiliser EN PRIORITÉ les informations des PIÈCES UPLOADÉES** (extraits fournis ci-dessous)
2. Le **formulaire** sert UNIQUEMENT de **FALLBACK** si l'info est absente des pièces
3. Si une entreprise/date/coordonnée est trouvée dans les pièces, ajoute la en plus des entreprises inscrites dans le formulaire sauf si tu reconnais le même nom d'entreprises.
**Structure & Style**
- Respecte les titres/numéros EXACTS de la trame (ne change rien)
- Si une info est introuvable : laisse vide
- Style : factuel, professionnel, concis
- Pas d'inventions : utilise uniquement les données réelles

**TABLEAUX — FORMAT CSV OBLIGATOIRE**
Pour chaque tableau, utilise EXACTEMENT ce format :

[TABLE:<Nom de section>]
"Col1";"Col2";"Col3"
"données";"données";"données"
[/TABLE]

- Point-virgule `;` comme séparateur
- Guillemets `"` autour de chaque cellule
- Si info manquante : laisse vide

Tableaux requis (en-têtes EXACTS) :
  [TABLE:1.3 Acteurs & coordonnées] "Acteur";"Société";"Nom";"Email";"Téléphone" [/TABLE]
  [TABLE:1.4 Planning] "Phase";"Activité";"Pré-requis";"Début";"Fin";"Responsable" [/TABLE]
  [TABLE:1.5 Effectifs] "Corps d'état";"Effectif max";"Habilitations";"Période" [/TABLE]
  [TABLE:1.6 Sous-traitants] "Entreprise";"Lot";"Responsable";"Contact";"Effectif";"Période";"Docs CSPS" [/TABLE]
  [TABLE:1.6 Matériel] "Matériel";"Caractéristiques";"VGP/Docs";"Responsable";"Période" [/TABLE]
  [TABLE:4. Prévention / EPI / Risques] "Risque";"Gravité";"Probabilité";"Criticité initiale";"Mesures";"Criticité résiduelle" [/TABLE]
  [TABLE:5. Secours & évacuation] "Rôle";"Qui";"Contact";"Back-up" [/TABLE]
  [TABLE:6. Mise à jour / révisions du PPSPS] "Index";"Date";"Motif";"Sections impactées";"Diffusion effectuée" [/TABLE]

**INSERTION D'IMAGES — RÈGLES STRICTES**

✅ **Images dans le CORPS du document** :
- Section "Secours & évacuation" : 1 plan évacuation [cat=securite_secours]
  Format : **Figure 1 — Plan d'évacuation Zone X**
  Légende : Points de rassemblement, issues, DAE | Source: [fichier] | Date: [si dispo]
  [IMAGE:nom_fichier.png]

- Section "Organisation / Circulation" : 1 PICH [cat=plan_circulation]
  Format : **Figure 2 — Plan d'installation de chantier (PICH)**
  Légende : Accès, zones stockage, balisage | Source: [fichier]
  [IMAGE:nom_fichier.png]
  _(Voir Annexe A pour plans détaillés)_

- Section "Modes opératoires / Levage" : 1 schéma [cat=plan_levage]
  Format : **Figure 3 — Schéma de levage / Zone de grutage**
  Légende : Rayons, zones interdites | Source: [fichier]
  [IMAGE:nom_fichier.png]
  _(Voir Annexe B pour détails)_

- Section "Réseaux" : PAS D'IMAGE, seulement texte :
  "Les prescriptions DICT/DT ont été intégrées. Voir plans réseaux en Annexe C."

- Section "Risques chimiques" : PAS D'IMAGE, seulement texte :
  "Produits employés : [liste]. Voir FDS complètes en Annexe D."

📎 **ANNEXES (À LA FIN DU DOCUMENT)** :

# Annexes

## Annexe A — Plans de circulation et PICH

**IMPORTANT** : Utilise UNIQUEMENT les noms de fichiers du catalogue fourni (section IMAGE_CATALOG).
Si aucune image [cat=plan_circulation] disponible, écrire : "Aucun document fourni"

Exemple de format SI des images existent :
[IMAGE:nom_reel_du_fichier.png]
**Plan PICH — Vue générale**
Source: [nom fichier] | Date: [si dispo]

## Annexe B — Plans de levage et manutention

Si des images [cat=plan_levage] existent dans le catalogue, les insérer.
Sinon : "Aucun document fourni"

## Annexe C — Plans de réseaux (DICT/DT)

Si des images [cat=plan_reseaux] existent dans le catalogue, les insérer.
Sinon : "Aucun document fourni"

## Annexe D — Fiches de Données de Sécurité (FDS)

Si des images [cat=chimique_fds] existent dans le catalogue, les insérer.
Sinon : "Aucun document fourni"

## Annexe E — Documents complémentaires

Laisser vide ou mentionner si VGP/échafaudages fournis.

🚨 **RÈGLES CRITIQUES IMAGES** :
- N'utilise QUE les noms de fichiers présents dans IMAGE_CATALOG fourni ci-dessous
- NE PAS inventer de noms comme "pich_1.png", "levage_1.png" etc.
- Si IMAGE_CATALOG est vide ou n'a pas d'images pour une catégorie → écrire "Aucun document fourni"
- Format exact : [IMAGE:nom_exact_du_fichier_du_catalogue.png]

**Modes opératoires**
- Si MOs détectés dans extraits : les réécrire proprement (Étapes, EPI, Prévention, Points de contrôle)
- Sinon, les déduire des travaux à réaliser
"""



    messages = [
        {"role": "system", "content": "Tu es un expert prévention SPS qui rédige des PPSPS conformes et rigoureux."},
        {"role": "user", "content":
            f"{RULES}\n\n"
            f"### TRAME À RESPECTER (NE RIEN MODIFIER)\n{TRAME_MD}\n\n"
            f"### 📄 PIÈCES UPLOADÉES (PRIORITÉ ABSOLUE)\n{evidence_pack}\n\n"
            f"### 📝 FORMULAIRE (FALLBACK UNIQUEMENT)\n{json.dumps(meta_hint, ensure_ascii=False, indent=2)}\n\n"
            f"### 🖼️ CATALOGUE D'IMAGES DISPONIBLES\n{json.dumps(img_catalog, ensure_ascii=False, indent=2)}\n\n"
            f"Maintenant, génère le PPSPS complet en respectant STRICTEMENT toutes les règles ci-dessus."
            f"\n\n### IMAGE_CATALOG (images disponibles)\n"
            f"{json.dumps(img_catalog, ensure_ascii=False)}\n"
            f"\n\n### RÈGLES D’INSERTION D’IMAGES (Mise à jour complète)\n"
            f"Insertion d’images (OBLIGATOIRE SI DISPONIBLE) :\n"
            f"- \"Organisation des secours\" : insérer le plan d’évacuation [cat=securite_secours], titré \"Figure X — Plan d’évacuation (zone)\", légende (points de rassemblement/cheminements), source (fichier) et date/vers. Utiliser [IMAGE:<filename>].\n"
            f"- \"Installations de chantier / Circulation\" : insérer le PICH principal [cat=plan_circulation], même format de titrage ; si plusieurs, n’en garder qu’1 (la plus lisible). Ajouter \"(Voir Annexe A pour détails)\".\n"
            f"- \"Modes opératoires / Levage & manutention\" : insérer 1 visuel pertinent [cat=plan_levage], titré. Ajouter \"(Voir Annexe B pour détails)\".\n"
            f"- \"Réseaux / DICT\" : ne pas insérer d’image ; créer un paragraphe qui dit \"Prescriptions DICT prises en compte — Voir Annexe C\".\n"
            f"- \"Risques chimiques\" : lister les produits employés + \"Voir Annexe D\" (ne pas coller de scans FDS en corps).\n\n"
            f"Titrage standard :\n"
            f"- Chaque image dans le corps : \"Figure X — <Titre>\", + légende courte (zone/objet), + \"Source: <fichier>\", + \"Date/Version\" si dispo.\n\n"
            f"Fallback :\n"
            f"- Si aucune image pertinente pour une section clé (évacuation, circulation, levage) : insérer un encart \"⚠️ Plan manquant — à fournir\".\n"
            f"- Plan de levage / manutention → **Annexe B**.\n"
            f"- Plans réseaux / DICT/DT → **Annexe C**.\n"
            f"- Pictos/FDS produits → **Annexe D**.\n"
            f"- VGP/échafaudages → **Annexe E**.\n"
            f"- Si une image semble utile en **Secours & évacuation** (point de rassemblement/DAE), insère-la dans cette section.\n"
            f"- Utilise **strictement** la balise `[IMAGE:<file>]` fournie dans le catalogue (ne pas renommer).\n"

        }
    ]
    return messages

def _delete_project_storage(project_id: int):
    base = _images_outdir(project_id)  # uploads/<id>/images
    root = _project_upload_dir(project_id)  # uploads/<id>
    for d in [base, root]:
        try:
            d_abs = _safe_path(UPLOADS_ROOT, os.path.relpath(d, start=UPLOADS_ROOT))
            if os.path.exists(d_abs):
                import shutil
                shutil.rmtree(d_abs, ignore_errors=True)
        except Exception:
            pass


# =====================================================================
#                   VALIDATION LÉGÈRE DU MARKDOWN
# =====================================================================

def _ensure_sections(md: str) -> str:
    """
    S'assure que toutes les sections de TRAME_TITLES existent. Ajoute le titre si manquant.
    """
    out = md or ""
    for title in TRAME_TITLES:
        h = title
        if not re.search(re.escape(title.split(" ", 1)[1]) if ". " in title else re.escape(title), out, flags=re.IGNORECASE):
            # inject just the heading
            if title.startswith("1. "): 
                out += f"\n\n# {title}\n"
            elif title.startswith(("2.", "3.", "4.", "5.", "6.", "7.")) and len(title) <= 4+len(title.split('.',1)[1]):
                out += f"\n\n# {title}\n"
            else:
                out += f"\n\n## {title}\n"
    return out

def _normalize_md(md: str) -> str:
    md = (md or "").strip()
    # Remplace les fences par leur contenu (évite de casser le parseur table)
    md = re.sub(r"```(.*?)```", lambda m: "\n" + m.group(1) + "\n", md, flags=re.DOTALL)
    return md

def _validate_and_fix_markdown(md: str) -> str:
    md = _normalize_md(md)
    md = _ensure_sections(md)
    return md

# === Tables attendues (ordre & colonnes exactes)
EXPECTED_TABLES = {
    "1.3 Acteurs & coordonnées": ["Acteur", "Société", "Nom", "Email", "Téléphone"],
    "1.4 Planning": ["Phase", "Activité", "Pré-requis", "Début", "Fin", "Responsable"],
    "1.5 Effectifs": ["Corps d’état", "Effectif max", "Habilitations", "Période"],
    "1.6 Sous-traitants": ["Entreprise", "Lot", "Responsable", "Contact", "Effectif", "Période", "Docs CSPS"],
    "1.6 Matériel": ["Matériel", "Caractéristiques", "VGP/Docs", "Responsable", "Période"],
    "4. Prévention / EPI / Risques": ["Risque", "Gravité", "Probabilité", "Criticité initiale", "Mesures", "Criticité résiduelle"],
    "5. Secours & évacuation": ["Rôle", "Qui", "Contact", "Back-up"],
    "6. Mise à jour / révisions du PPSPS": ["Index", "Date", "Motif", "Sections impactées", "Diffusion effectuée"],
}

TABLE_BLOCK_RE = re.compile(r"\[TABLE:([^\]]+)\](.*?)\[/TABLE\]", re.DOTALL)
IMAGE_BLOCK_RE = re.compile(r"\[IMAGE:([^\]]+)\]")


def _split_text_and_tables(md: str):
    """
    Découpe le markdown en segments:
      [("text", "..."), ("table", section_name, csv_text), ...]
    """
    segs = []
    pos = 0
    for m in TABLE_BLOCK_RE.finditer(md or ""):
        start, end = m.span()
        if start > pos:
            segs.append(("text", md[pos:start]))
        section = m.group(1).strip()
        csv_text = m.group(2).strip()
        segs.append(("table", section, csv_text))
        pos = end
    if pos < len(md or ""):
        segs.append(("text", md[pos:]))
    return segs

def _parse_csv_table(csv_text: str) -> list[list[str]]:
    """
    Parse CSV (séparateur ';', guillemets obligatoires) → rows[list[str]]
    Tolère lignes vides; strip global.
    """
    rows = []
    reader = csv.reader([l for l in csv_text.splitlines()], delimiter=';', quotechar='"')
    for r in reader:
        if not r: 
            continue
        rows.append([c.strip() for c in r])
    return rows

def _ensure_expected_header(section: str, rows: list[list[str]]) -> list[list[str]]:
    """
    Vérifie l'en-tête; si KO, impose l'en-tête attendu et une ligne vide.
    """
    expected = EXPECTED_TABLES.get(section)
    if not expected:
        return rows
    if not rows:
        return [expected, [""]*(len(expected))]
    header = rows[0]
    if header != expected:
        return [expected, [""]*(len(expected))]
    return rows

def _docx_add_csv_table(doc, section: str, csv_text: str):
    rows = _parse_csv_table(csv_text)
    rows = _ensure_expected_header(section, rows)
    if not rows:
        return
    # crée table Word
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Table Grid"
    for i, r in enumerate(rows):
        for j, cell in enumerate(r):
            tbl.cell(i, j).text = cell
    _format_table_pretty(tbl, header_fill=True)



# =====================================================================
#            ENDPOINT PRINCIPAL : FREEFORM PPSPS (Markdown)
# =====================================================================

@app.post("/projects/{project_id}/generate_ppsps_freeform")
def generate_ppsps_freeform(project_id: int, session: Session = Depends(get_session), 
                           user: UserDB = Depends(require_login)):
    """
    Génère un PPSPS en utilisant le template DOCX et en le remplissant intelligemment avec l'IA.
    Sauvegarde le DOCX généré pour export ultérieur.
    """
    proj = session.get(ProjectDB, project_id)
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # ✅ VÉRIFIER ET CONSOMMER 1 JETON
    try:
        TokenService.use_token(
            session=session,
            user_id=user.id,
            project_id=project_id,
            description=f"Génération PPSPS - {proj.name}"
        )
    except InsufficientTokensError as e:
        raise HTTPException(
            status_code=402,
            detail={
                "error": "insufficient_tokens",
                "message": str(e),
                "redirect": "/tokens/shop"
            }
        )
    
    # ✅ VÉRIFIER LA PERTINENCE DES FICHIERS AVEC L'IA
    is_relevant, relevance_message = _check_files_relevance_with_ai(session, project_id)

    if not is_relevant:
        TokenService.refund_token(
            session=session,
            user_id=user.id,
            project_id=project_id,
            reason="Documents non pertinents - remboursement automatique"
        )
        raise HTTPException(status_code=400, detail=relevance_message)

    if relevance_message:
        logger.warning(f"[PPSPS] {relevance_message}")
    
    # Récupération des images
    all_images = session.exec(
        select(AttachmentDB).where(
            AttachmentDB.project_id == project_id,
            AttachmentDB.mime_type.like("image%")
        ).order_by(AttachmentDB.created_at.asc())
    ).all()
    
    img_catalog = [
        {
            "file": img.filename,
            "stored_path": img.stored_path,
            "size": img.size_bytes
        }
        for img in all_images
    ]
    
    logger.info(f"[IMG] {len(img_catalog)} images disponibles pour le projet {project_id}")

    # Evidence depuis les pièces
    blob = _project_text_blob(session, project_id, limit_chars=80_000)
    evidence = _build_evidence_pack(blob, proj)
    meta_hint = _build_meta_hint(proj)

    # Vérifier que le template existe
    if not os.path.exists(TEMPLATE_PATH):
        TokenService.refund_token(
            session=session,
            user_id=user.id,
            project_id=project_id,
            reason="Template PPSPS introuvable - remboursement automatique"
        )
        raise HTTPException(status_code=500, detail="Template PPSPS introuvable")
    
    # Utiliser le TemplateFiller pour remplir le template
    try:
        # Préparer les données du formulaire pour les placeholders
        form_data = {
            "name": proj.name or "",
            "address": proj.address or "",
            "project_reference": getattr(proj, 'project_reference', "") or "",
            "site_phone": getattr(proj, 'site_phone', "") or "",
            "duration_weeks": proj.duration_weeks or 0,
            "workforce": proj.workforce or 0,
            "works_csv": proj.works_csv or "",
            "company_name": getattr(proj, 'company_name', "") or "",
            "company_address": getattr(proj, 'company_address', "") or "",
            "company_phone": getattr(proj, 'company_phone', "") or "",
            "company_email": getattr(proj, 'company_email', "") or "",
            "site_manager_name": getattr(proj, 'site_manager_name', "") or "",
            "owner_name": getattr(proj, 'owner_name', "") or "",
            "architect_name": getattr(proj, 'architect_name', "") or "",
        }
        
        filler = TemplateFiller(TEMPLATE_PATH, form_data=form_data)
        filled_doc = filler.fill_with_ai(
            project_data=meta_hint,
            evidence_pack=evidence,
            img_catalog=img_catalog,
            openai_client=client,
            model=OPENROUTER_DEFAULT_MODEL
        )
        
        # Déterminer la version
        existing = session.exec(
            select(DocumentDB).where(
                DocumentDB.project_id == project_id,
                DocumentDB.doc_type == "PPSPS"
            )
        ).all()
        next_version = (max([d.version for d in existing]) + 1) if existing else 1
        
        # Sauvegarder le DOCX généré dans un fichier
        project_dir = _safe_path(UPLOADS_ROOT, _project_upload_dir(project_id))
        os.makedirs(project_dir, exist_ok=True)
        docx_filename = f"PPSPS_v{next_version}.docx"
        docx_path = _safe_path(project_dir, docx_filename)
        filled_doc.save(docx_path)
        
        logger.info(f"[PPSPS] DOCX sauvegardé : {docx_path}")
        
        # Enregistrer dans la base de données avec le chemin du DOCX
        dbdoc = DocumentDB(
            project_id=project_id,
            doc_type="PPSPS",
            content_md=f"[DOCX généré : {docx_path}]",
            version=next_version,
            status="draft",
        )
        session.add(dbdoc)
        session.commit()
        session.refresh(dbdoc)
        
        logger.info(f"[PPSPS] Document DB créé (ID: {dbdoc.id}, version {next_version})")
        
        return {
            "ok": True,
            "document_id": dbdoc.id,
            "version": dbdoc.version,
            "docx_path": docx_path,
            "message": "PPSPS généré avec succès depuis le template"
        }
        
    except json.JSONDecodeError as e:
        logger.error(f"[PPSPS] Erreur parsing JSON de l'IA : {e}")
        TokenService.refund_token(
            session=session,
            user_id=user.id,
            project_id=project_id,
            reason="Erreur de génération (JSON invalide) - remboursement automatique"
        )
        raise HTTPException(status_code=502, detail="Erreur de génération : réponse IA invalide")
    except Exception as e:
        logger.error(f"[PPSPS] Erreur génération : {e}")
        TokenService.refund_token(
            session=session,
            user_id=user.id,
            project_id=project_id,
            reason="Erreur de génération - remboursement automatique"
        )
        raise HTTPException(status_code=502, detail=f"Génération indisponible : {str(e)}")




@app.post("/projects/{project_id}/files")
def upload_file(project_id: int, file: UploadFile = File(...), session: Session = Depends(get_session), user: UserDB = Depends(require_login)):
    proj = get_owned_project_or_404(project_id, user, session)
    original = _safe_name(file.filename or "file")
    ext = _ext(original)
    if ext not in ALLOWED_EXTS:
        raise HTTPException(status_code=400, detail=f"Extension non autorisée ({ext}).")
    contents = file.file.read()
    size = len(contents)
    if size > MAX_FILE_MB * 1024 * 1024:
        raise HTTPException(status_code=400, detail=f"Fichier trop volumineux (> {MAX_FILE_MB} Mo).")

    uid = datetime.now().strftime("%Y%m%d%H%M%S%f")
    stored_name = f"{uid}_{original}"
    folder = _project_upload_dir(proj.id)
    stored_path = os.path.join(folder, stored_name)
    with open(stored_path, "wb") as out:
        out.write(contents)

    extracted = extract_text_from_file(stored_path)

    # extraction d'images associées (inchangé)
    img_dir = _images_outdir(proj.id)
    new_imgs = []
    if ext == ".docx":
        for fp in _extract_images_from_docx(stored_path, img_dir):
            new_imgs.append((fp, None))
    elif ext == ".pdf":
        try:
            keep = peek_pages_for_plans(stored_path)
        except Exception:
            keep = None
        for fp, page_no in _extract_images_from_pdf(stored_path, img_dir, keep_pages=keep):
            new_imgs.append((fp, page_no))

    for fp, page_no in new_imgs:
        att_img = AttachmentDB(
            project_id=proj.id,
            filename=os.path.basename(fp),
            stored_path=fp,
            mime_type="image/png",
            size_bytes=os.path.getsize(fp),
            extracted_text=f"(image extraite de {original} page {page_no})" if page_no else "(image extraite)",
        )
        session.add(att_img)
    session.commit()

    att = AttachmentDB(
        project_id=proj.id,
        filename=original,
        stored_path=stored_path,
        mime_type=file.content_type or "application/octet-stream",
        size_bytes=size,
        extracted_text=extracted[:500_000] if extracted else None,
    )
    session.add(att)
    session.commit()
    session.refresh(att)
    return att

@app.get("/projects/{project_id}/files")
def list_files(project_id: int, session: Session = Depends(get_session), user: UserDB = Depends(require_login)):
    proj = get_owned_project_or_404(project_id, user, session)
    rows = session.exec(
        select(AttachmentDB).where(AttachmentDB.project_id == proj.id).order_by(AttachmentDB.created_at.desc())
    ).all()
    return rows

@app.get("/files/{file_id}/download")
def download_file(file_id: int, session: Session = Depends(get_session), user: UserDB = Depends(require_login)):
    att = ensure_file_is_owned(file_id, user, session)
    # Assure que le chemin reste sous uploads/
    stored = _safe_path(UPLOADS_ROOT, os.path.relpath(att.stored_path, start=UPLOADS_ROOT))
    if not os.path.exists(stored):
        raise HTTPException(status_code=404, detail="Fichier introuvable")
    return FileResponse(stored, media_type=att.mime_type, filename=att.filename)

from fastapi import Header

@app.delete("/files/{file_id}")
def delete_file(
    file_id: int,
    request: Request,
    x_csrf_token: str | None = Header(None, alias="X-CSRF-Token"),
    session: Session = Depends(get_session),
    user: UserDB = Depends(require_login),
):
    _check_csrf(request, x_csrf_token)
    att = ensure_file_is_owned(file_id, user, session)
    try:
        if os.path.exists(att.stored_path):
            os.remove(att.stored_path)
    except:
        pass
    session.delete(att)
    session.commit()
    return {"ok": True}

def _image_lookup_for_project(session: Session, project_id: int) -> dict[str, str]:
    rows = session.exec(
        select(AttachmentDB).where(
            AttachmentDB.project_id == project_id,
            AttachmentDB.mime_type.like("image%")
        )
    ).all()
    return { (r.filename or "").strip(): (r.stored_path or "").strip() for r in rows if r.filename and r.stored_path }

def get_owned_project_or_404(project_id: int, user: UserDB, session: Session) -> ProjectDB:
    proj = session.get(ProjectDB, project_id)
    if not proj or getattr(proj, "owner_id", None) != user.id:
        # on renvoie 404 pour ne rien révéler
        raise HTTPException(status_code=404, detail="Project not found")
    return proj

def ensure_doc_is_owned(doc_id: int, user: UserDB, session: Session) -> DocumentDB:
    doc = session.get(DocumentDB, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    proj = session.get(ProjectDB, doc.project_id)
    if not proj or getattr(proj, "owner_id", None) != user.id:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc

def ensure_file_is_owned(file_id: int, user: UserDB, session: Session) -> AttachmentDB:
    att = session.get(AttachmentDB, file_id)
    if not att:
        raise HTTPException(status_code=404, detail="Fichier introuvable")
    proj = session.get(ProjectDB, att.project_id)
    if not proj or getattr(proj, "owner_id", None) != user.id:
        raise HTTPException(status_code=404, detail="Fichier introuvable")
    return att


# =====================================================================
#                           ROUTES SEO
# =====================================================================
@app.get("/sitemap.xml")
def sitemap():
    urls = [
        {"loc": SEOConfig.SITE_URL, "priority": "1.0", "changefreq": "daily"},
        {"loc": f"{SEOConfig.SITE_URL}/home", "priority": "1.0", "changefreq": "daily"},
        {"loc": f"{SEOConfig.SITE_URL}/register", "priority": "0.8", "changefreq": "monthly"},
        {"loc": f"{SEOConfig.SITE_URL}/tokens/shop", "priority": "0.9", "changefreq": "weekly"},
    ]
    
    sitemap_xml = '<?xml version="1.0" encoding="UTF-8"?>\n'
    sitemap_xml += '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n'
    for url in urls:
        sitemap_xml += f'  <url>\n'
        sitemap_xml += f'    <loc>{url["loc"]}</loc>\n'
        sitemap_xml += f'    <priority>{url["priority"]}</priority>\n'
        sitemap_xml += f'    <changefreq>{url["changefreq"]}</changefreq>\n'
        sitemap_xml += f'  </url>\n'
    sitemap_xml += '</urlset>'
    
    from fastapi.responses import Response
    return Response(content=sitemap_xml, media_type="application/xml")

@app.get("/robots.txt")
def robots_txt():
    robots = f"""User-agent: *
Allow: /
Disallow: /ui/
Disallow: /api/

Sitemap: {SEOConfig.SITE_URL}/sitemap.xml
"""
    from fastapi.responses import Response
    return Response(content=robots, media_type="text/plain")
@app.get("/", include_in_schema=False)
async def redirect_root(request: Request):
    # Si l'utilisateur est connecté, aller vers /ui/projects
    if request.session.get("uid"):
        return RedirectResponse(url="/ui/projects")
    # Sinon montrer la homepage
    return RedirectResponse(url="/home")


@app.get("/home", response_class=HTMLResponse)
def home_page(request: Request):
    """Page d'accueil avec contenu SEO"""
    seo_config = SEO_PAGES.get("home", {})
    meta = SEOConfig.get_meta_tags(
        title=seo_config.get("title"),
        description=seo_config.get("description"),
        keywords=seo_config.get("keywords"),
        canonical_url=SEOConfig.SITE_URL
    )
    
    structured_data = {
        "@context": "https://schema.org",
        "@type": "SoftwareApplication",
        "name": "PPSPS Generator",
        "applicationCategory": "BusinessApplication",
        "operatingSystem": "Web",
        "offers": {
            "@type": "Offer",
            "price": "50",
            "priceCurrency": "EUR",
            "availability": "https://schema.org/InStock"
        },
        "aggregateRating": {
            "@type": "AggregateRating",
            "ratingValue": "4.8",
            "reviewCount": "127"
        },
        "description": "Générateur automatique de PPSPS par IA pour les professionnels du BTP"
    }
    
    return templates.TemplateResponse("index.html", {
        "request": request,
        "meta": meta,
        "structured_data": structured_data
    })



# =====================================================================
#                               UI (Jinja)
# =====================================================================

@app.get("/login", response_class=HTMLResponse)
def ui_login(request: Request):
    csrf = _ensure_csrf_token(request)
    return templates.TemplateResponse("login.html", {"request": request, "error": None, "csrf_token": csrf})

@app.post("/login")
def ui_login_post(
    request: Request,
    email: str = Form(...),
    password: str = Form(...),
    csrf_token: str = Form(...),
    session: Session = Depends(get_session),
):
    _check_csrf(request, csrf_token)
    user = session.exec(select(UserDB).where(UserDB.email == email.strip().lower())).first()
    if not user or not verify_pwd(password, user.password_hash):
        csrf = _ensure_csrf_token(request)
        return templates.TemplateResponse("login.html", {"request": request, "error": "Email ou mot de passe invalide.", "csrf_token": csrf}, status_code=400)
    request.session["uid"] = user.id
    return RedirectResponse(url="/ui/projects", status_code=303)

@app.post("/logout")
def ui_logout(request: Request):
    request.session.clear()
    return RedirectResponse(url="/login", status_code=303)

@app.get("/register", response_class=HTMLResponse)
def ui_register(request: Request):
    csrf = _ensure_csrf_token(request)
    return templates.TemplateResponse("register.html", {"request": request, "error": None, "csrf_token": csrf})

@app.post("/register")
def ui_register_post(
    request: Request,
    email: str = Form(...),
    password: str = Form(...),
    csrf_token: str = Form(...),
    session: Session = Depends(get_session),
):
    _check_csrf(request, csrf_token)
    email_n = email.strip().lower()

    def strong(p: str) -> bool:
        return (
            len(p) >= 8 and
            any(c.islower() for c in p) and
            any(c.isupper() for c in p) and
            any(c.isdigit() for c in p)
        )

    if not email_n or not strong(password):
        csrf = _ensure_csrf_token(request)
        return templates.TemplateResponse(
            "register.html",
            {
                "request": request,
                "error": "Email valide et mot de passe robuste requis (≥8, majuscule, minuscule, chiffre).",
                "csrf_token": csrf
            },
            status_code=400
        )
    exists = session.exec(select(UserDB).where(UserDB.email == email_n)).first()
    if exists:
        csrf = _ensure_csrf_token(request)
        return templates.TemplateResponse("register.html", {"request": request, "error": "Email déjà utilisé.", "csrf_token": csrf}, status_code=400)

    user = UserDB(email=email_n, password_hash=hash_pwd(password), is_admin=False)
    session.add(user); session.commit(); session.refresh(user)
    request.session["uid"] = user.id
    return RedirectResponse(url="/ui/projects", status_code=303)

@app.get("/ui/projects", response_class=HTMLResponse)
def ui_projects(request: Request, session: Session = Depends(get_session), user: UserDB = Depends(require_login)):
    projects = session.exec(select(ProjectDB).where(ProjectDB.owner_id == user.id)).all()
    csrf = _ensure_csrf_token(request)
    return templates.TemplateResponse("projects_list.html", {"request": request, "projects": projects, "csrf_token": csrf})

@app.post("/ui/projects/create")
def ui_create_project(
    request: Request,
    name: str = Form(...),
    address: str = Form(...),
    works: str = Form(""),
    duration_weeks: int = Form(0),
    workforce: int = Form(0),
    companies: str = Form(""),
    project_reference: str = Form(""),
    site_phone: str = Form(""),
    company_name: str = Form(""),
    company_address: str = Form(""),
    company_phone: str = Form(""),
    company_email: str = Form(""),
    site_manager_name: str = Form(""),
    owner_name: str = Form(""),
    architect_name: str = Form(""),
    csrf_token: str = Form(...),
    session: Session = Depends(get_session),
    user: UserDB = Depends(require_login),
):
    _check_csrf(request, csrf_token)

    pr = ProjectDB(
        name=name.strip(),
        address=address.strip(),
        works_csv=",".join([w.strip() for w in works.split(",") if w.strip()]),
        duration_weeks=duration_weeks or 0,
        workforce=workforce or 0,
        companies_csv=",".join([c.strip() for c in companies.split(",") if c.strip()]),
        project_reference=project_reference.strip(),
        site_phone=site_phone.strip(),
        company_name=company_name.strip(),
        company_address=company_address.strip(),
        company_phone=company_phone.strip(),
        company_email=company_email.strip(),
        site_manager_name=site_manager_name.strip(),
        owner_name=owner_name.strip(),
        architect_name=architect_name.strip(),
        owner_id=user.id,
    )
    session.add(pr)
    session.commit()
    session.refresh(pr)
    return RedirectResponse(url=f"/ui/projects/{pr.id}", status_code=303)

@app.get("/ui/projects/{project_id}", response_class=HTMLResponse)
def ui_project_detail(project_id: int, request: Request, session: Session = Depends(get_session), user: UserDB = Depends(require_login)):
    proj = get_owned_project_or_404(project_id, user, session)
    docs = session.exec(
        select(DocumentDB).where(DocumentDB.project_id == proj.id).order_by(DocumentDB.created_at.desc())
    ).all()
    files = session.exec(
    select(AttachmentDB)
    .where(
        AttachmentDB.project_id == proj.id,
        ~AttachmentDB.extracted_text.like("(image extraite%"))
    .order_by(AttachmentDB.created_at.desc())
    ).all()
    ia_summary = {"mapped_fields": 0, "taches": 0, "sous_traitants": 0}
    csrf = _ensure_csrf_token(request)
    # Récupérer le solde de jetons
    balance = TokenService.get_balance(session, user.id)
    return templates.TemplateResponse("project_detail.html", {
        "request": request, "p": proj, "docs": docs, "files": files, "ia_summary": ia_summary, "csrf_token": csrf, "balance": balance
    })

@app.post("/ui/projects/{project_id}/ingest")
def ui_project_ingest(project_id: int, csrf_token: str = Form(...), request: Request = None, session: Session = Depends(get_session), user: UserDB = Depends(require_login)):
    _check_csrf(request, csrf_token)
    proj = get_owned_project_or_404(project_id, user, session)
    _ = ingest_project_files(proj.id, session)
    return RedirectResponse(url=f"/ui/projects/{proj.id}", status_code=303)

@app.post("/ui/projects/{project_id}/files/upload")
async def ui_upload_file(
    project_id: int,
    request: Request,
    file: UploadFile = File(...),
    csrf_token: str = Form(...),
    session: Session = Depends(get_session),
    user: UserDB = Depends(require_login),
):
    _check_csrf(request, csrf_token)
    proj = get_owned_project_or_404(project_id, user, session)

    # Quotas projet (exclure images extraites)
    MAX_FILES = 50
    MAX_TOTAL_MB = 200
    existing = session.exec(
        select(AttachmentDB).where(
            AttachmentDB.project_id == proj.id,
            ~AttachmentDB.extracted_text.like("(image extraite%")
        )
    ).all()
    if len(existing) >= MAX_FILES:
        raise HTTPException(status_code=400, detail="Quota fichiers atteint (50).")
    total_size = sum((r.size_bytes or 0) for r in existing)

    # Lecture contrôlée
    contents = file.file.read()
    size = len(contents)
    if size > MAX_FILE_MB * 1024 * 1024:
        raise HTTPException(status_code=400, detail=f"Fichier trop volumineux (> {MAX_FILE_MB} Mo).")
    if (total_size + size) > (MAX_TOTAL_MB * 1024 * 1024):
        raise HTTPException(status_code=400, detail="Quota de stockage projet dépassé (200 Mo).")

    # Validation extension & MIME & signature
    original = _safe_name(file.filename or "file")
    ext = _ext(original)
    if ext not in ALLOWED_EXTS:
        raise HTTPException(status_code=400, detail=f"Extension non autorisée ({ext}).")

    mime = (file.content_type or "").lower()
    allowed_mimes = {
        ".pdf": {"application/pdf"},
        ".docx": {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/zip",
        },
        ".txt": {"text/plain"},
        ".md": {"text/markdown", "text/plain"},
    }
    # tolérance MIME client imprécis, on vérifiera la signature
    header = contents[:8]

    def _looks_pdf(h: bytes) -> bool: return h.startswith(b"%PDF-")
    def _looks_docx(b: bytes) -> bool: return b[:2] == b"PK"
    def _looks_text(b: bytes) -> bool:
        try:
            b.decode("utf-8"); return True
        except Exception:
            return False

    if ext == ".pdf" and not _looks_pdf(header):
        raise HTTPException(status_code=400, detail="Signature PDF invalide.")
    if ext == ".docx" and not _looks_docx(contents):
        raise HTTPException(status_code=400, detail="Signature DOCX invalide (ZIP attendu).")
    if ext in (".txt", ".md") and not _looks_text(contents[:2048]):
        raise HTTPException(status_code=400, detail="Contenu texte invalide.")

    # Écriture disque (chemin sûr)
    uid = datetime.now().strftime("%Y%m%d%H%M%S%f")
    stored_name = f"{uid}_{original}"
    folder = _safe_path(UPLOADS_ROOT, _project_upload_dir(proj.id))
    os.makedirs(folder, exist_ok=True)
    stored_path = _safe_path(folder, stored_name)
    with open(stored_path, "wb") as out:
        out.write(contents)

    # Extraction texte
    extracted = extract_text_from_file(stored_path)

    # Extraction images (PDF/DOCX) -> dans img_dir
    img_dir = _safe_path(UPLOADS_ROOT, _images_outdir(proj.id))
    os.makedirs(img_dir, exist_ok=True)

    new_imgs: list[tuple[str, int | None]] = []
    if ext == ".docx":
        for fp in _extract_images_from_docx(stored_path, img_dir):
            new_imgs.append((fp, None))
    elif ext == ".pdf":
        try:
            keep = peek_pages_for_plans(stored_path)
        except Exception:
            keep = None
        for fp, page_no in _extract_images_from_pdf(stored_path, img_dir, keep_pages=keep):
            new_imgs.append((fp, page_no))

    # === PATCH 2 : nom unique + chemin réel et sûr pour CHAQUE image extraite ===
    import uuid, time, shutil
    for fp, page_no in new_imgs:
        base = os.path.basename(fp)
        _, ext_img = os.path.splitext(base)
        ext_img = (ext_img or ".png").lower()
        uniq = f"{int(time.time())}_{uuid.uuid4().hex[:8]}{ext_img}"
        dst = _safe_path(img_dir, uniq)

        # si l'extract a déjà mis le fichier dans img_dir, on renomme proprement
        try:
            shutil.move(fp, dst) if os.path.dirname(fp) == img_dir else shutil.copy2(fp, dst)
        except Exception:
            # fallback copie
            shutil.copy2(fp, dst)

        att_img = AttachmentDB(
            project_id=proj.id,
            filename=base,
            stored_path=dst,  # chemin FINAL réel
            mime_type="image/png",
            size_bytes=os.path.getsize(dst),
            extracted_text=(f"(image extraite de {original} page {page_no})" if page_no else "(image extraite)"),
        )
        session.add(att_img)

    session.commit()

    # Pièce jointe principale (le fichier uploadé)
    att = AttachmentDB(
        project_id=proj.id,
        filename=original,
        stored_path=stored_path,
        mime_type=file.content_type or "application/octet-stream",
        size_bytes=size,
        extracted_text=extracted[:500_000] if extracted else None,
    )
    session.add(att)
    session.commit()
    session.refresh(att)

    return RedirectResponse(url=f"/ui/projects/{proj.id}", status_code=303)

@app.api_route("/ui/projects/{project_id}/gen/{kind}", methods=["POST"])
async def ui_generate_doc(
    project_id: int,
    kind: str,
    request: Request,
    session: Session = Depends(get_session),
    user: UserDB = Depends(require_login),
):
    form = await request.form()
    token = (form.get("csrf_token"))
    _check_csrf(request, token)

    # 1) récupérer le projet d'abord
    proj = get_owned_project_or_404(project_id, user, session)

    # 2) refuser la génération s'il n'y a aucun fichier
    has_files = session.exec(
        select(AttachmentDB.id).where(AttachmentDB.project_id == proj.id)
    ).first()
    if not has_files:
        raise HTTPException(
            status_code=400,
            detail="Ajoutez au moins une pièce avant de générer le PPSPS."
        )

    # 3) dispatcher par type
    k = kind.lower()
    if k in ("ppsps", "ppsps_freeform", "ppsps_docx"):
        try:
            res = generate_ppsps_freeform(proj.id, session, user)
            doc_id = res.get("document_id")
            if not doc_id:
                raise HTTPException(status_code=500, detail="Génération PPSPS échouée")
            return export_docx_by_id(doc_id, session=session, user=user)
        except HTTPException as e:
            # Si c'est une erreur 402 (jetons insuffisants), renvoyer du JSON
            if e.status_code == 402:
                from fastapi.responses import JSONResponse
                return JSONResponse(
                    status_code=402,
                    content=e.detail
                )
            # Sinon, laisser l'exception se propager
            raise

    raise HTTPException(status_code=400, detail="Type de doc inconnu (attendu: ppsps)")

# =====================================================================
#                     SUPPRESSION Projet / Document
# =====================================================================
@app.delete("/projects/{project_id}")
def delete_project(
    project_id: int,
    request: Request,
    x_csrf_token: str | None = Header(None, alias="X-CSRF-Token"),
    session: Session = Depends(get_session),
    user: UserDB = Depends(require_login),
):
    _check_csrf(request, x_csrf_token)
    proj = session.get(ProjectDB, project_id)
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")

    docs = session.exec(select(DocumentDB).where(DocumentDB.project_id == project_id)).all()
    for d in docs:
        session.delete(d)

    atts = session.exec(select(AttachmentDB).where(AttachmentDB.project_id == project_id)).all()
    for a in atts:
        if a.stored_path and os.path.exists(a.stored_path):
            try: os.remove(a.stored_path)
            except: pass
        session.delete(a)

    session.delete(proj)
    session.commit()
    return {"ok": True}

@app.delete("/documents/{doc_id}")
def delete_document(
    doc_id: int,
    request: Request,
    x_csrf_token: str | None = Header(None, alias="X-CSRF-Token"),
    session: Session = Depends(get_session),
    user: UserDB = Depends(require_login),
):
    _check_csrf(request, x_csrf_token)
    doc = ensure_doc_is_owned(doc_id, user, session)
    session.delete(doc); session.commit()
    return {"ok": True}

# =====================================================================
#                           ROUTES TOKENS
# =====================================================================
@app.get("/tokens/balance")
def get_token_balance(session: Session = Depends(get_session), user: UserDB = Depends(require_login)):
    balance = TokenService.get_balance(session, user.id)
    return {"balance": balance}

@app.get("/tokens/shop", response_class=HTMLResponse)
def token_shop(request: Request, session: Session = Depends(get_session), user: UserDB = Depends(require_login)):
    packages = StripeService.get_packages(session)
    balance = TokenService.get_balance(session, user.id)
    transactions = TokenService.get_transactions(session, user.id, limit=10)
    
    seo_config = SEO_PAGES.get("shop", {})
    meta = SEOConfig.get_meta_tags(
        title=seo_config.get("title"),
        description=seo_config.get("description"),
        keywords=seo_config.get("keywords"),
        canonical_url=f"{SEOConfig.SITE_URL}/tokens/shop"
    )
    
    return templates.TemplateResponse("token_shop.html", {
        "request": request,
        "packages": packages,
        "balance": balance,
        "transactions": transactions,
        "user": user,
        "meta": meta
    })

@app.post("/tokens/purchase/{package_id}")
def purchase_tokens(package_id: int, request: Request, session: Session = Depends(get_session), 
                   user: UserDB = Depends(require_login)):
    base_url = str(request.base_url).rstrip('/')
    success_url = f"{base_url}/tokens/shop?success=true"
    cancel_url = f"{base_url}/tokens/shop?canceled=true"
    
    try:
        checkout_session = StripeService.create_checkout_session(
            user_id=user.id,
            package_id=package_id,
            success_url=success_url,
            cancel_url=cancel_url,
            session_db=session
        )
        return {"checkout_url": checkout_session.url}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception:
        raise HTTPException(status_code=500, detail="Erreur lors de la création du paiement")

from fastapi import Header

@app.post("/stripe/webhook")
async def stripe_webhook(request: Request, stripe_signature: str = Header(None, alias="stripe-signature"),
                        session: Session = Depends(get_session)):
    payload = await request.body()
    result = StripeService.handle_webhook(payload=payload, sig_header=stripe_signature, session_db=session)
    if not result['success']:
        raise HTTPException(status_code=400, detail=result['message'])
    return {"status": "success"}


# =====================================================================
#                         KB UI ADMIN (sécurisée)
# =====================================================================
KB_BASE = Path("app/kb").resolve()

def _safe_kb_path(rel_path: str) -> Path:
    rel_path = unquote(rel_path).strip().lstrip("/").replace("\\", "/")
    fs_path = (KB_BASE / rel_path).resolve()
    if KB_BASE not in fs_path.parents or not str(fs_path).endswith(".md"):
        raise HTTPException(status_code=404, detail="Fichier introuvable")
    return fs_path

@app.get("/ui/kb", response_class=HTMLResponse)
def ui_kb_list(request: Request):
    _require_admin(request)
    items = []
    for p in KB_BASE.rglob("*.md"):
        rel = p.relative_to(KB_BASE).as_posix()
        items.append({"rel": rel, "size": p.stat().st_size})
    items.sort(key=lambda x: x["rel"])
    return templates.TemplateResponse("kb_list.html", {"request": request, "items": items})

@app.get("/ui/kb/edit", response_class=HTMLResponse)
def ui_kb_edit(request: Request, path: str):
    _require_admin(request)
    fs_path = _safe_kb_path(path)
    if not fs_path.exists() or not fs_path.is_file():
        return HTMLResponse("Fichier introuvable", status_code=404)
    text = fs_path.read_text(encoding="utf-8", errors="replace")
    return templates.TemplateResponse("kb_edit.html", {"request": request, "path": path, "content": text})

@app.post("/ui/kb/save")
def ui_kb_save(request: Request, path: str = Form(...), content: str = Form(...)):
    _require_admin(request)
    rel_path = unquote(path).strip().lstrip("/").replace("\\", "/")
    fs_path = (KB_BASE / rel_path).resolve()
    if KB_BASE not in fs_path.parents or not str(fs_path).endswith(".md"):
        raise HTTPException(status_code=400, detail="Chemin non autorisé")
    fs_path.parent.mkdir(parents=True, exist_ok=True)
    fs_path.write_text(content, encoding="utf-8")
    return RedirectResponse(url=f"/ui/kb?token={ADMIN_TOKEN}", status_code=303)

# =====================================================================
#                            HEALTH / DEBUG
# =====================================================================
@app.get("/debug")
def debug():
    kb_files = glob.glob("app/kb/**/*.md", recursive=True)
    return {
        "kb_files_count": len(kb_files),
        "kb_files_sample": kb_files[:10],
    }

@app.get("/healthz")
def healthz():
    return {"ok": True}
